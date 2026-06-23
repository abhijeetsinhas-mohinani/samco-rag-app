import os
import io
import re
import json
import pickle
import struct
import logging
import time
import hashlib
import requests
import traceback
from difflib import SequenceMatcher
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

import numpy as np
import pandas as pd
from docx import Document as DocxDocument
import pdfplumber  # fallback for PDF
from docling_parse.pdf_parser import DoclingPdfParser  # primary PDF parser
from rank_bm25 import BM25Okapi
import faiss
import olefile
import mammoth
import subprocess
import tempfile
from azure.storage.blob import BlobServiceClient
from sentence_transformers import SentenceTransformer

# Docling DOCX backend — much better heading detection than plain python-docx
try:
    from docling.backend.msword_backend import MsWordDocumentBackend
    from docling.datamodel.document import InputDocument
    from docling.datamodel.base_models import InputFormat
    from docling_core.types.doc import DocItemLabel
    HAS_DOCLING_DOCX = True
except ImportError:
    HAS_DOCLING_DOCX = False

# ---------------------------------------------------------------------------
# Logging Setup
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("rag_pipeline")

# ---------------------------------------------------------------------------
# Configuration — loaded from .env / environment variables
# ---------------------------------------------------------------------------
AZURE_STORAGE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
DOCUMENT_CONTAINER             = os.getenv("DOCUMENT_CONTAINER", "prod-rag-documents")
VECTOR_CONTAINER               = os.getenv("VECTOR_CONTAINER", "vector-index")
OLLAMA_BASE_URL                = os.getenv("OLLAMA_BASE_URL", "http://4.210.114.244:11434")
MODEL_NAME                     = os.getenv("MODEL_NAME", "llama3:8b")
EMBEDDING_MODEL                = os.getenv("EMBEDDING_MODEL", "BAAI/bge-base-en-v1.5")
FAISS_LOCAL_PATH               = os.getenv("FAISS_LOCAL_PATH", "faiss_temp")
CHUNK_SIZE                     = int(os.getenv("CHUNK_SIZE", "350"))
# ↑ Matches BAAI/bge-base-en-v1.5 token limit (~512 tokens ≈ 350 words).
#   If you switch to a model with larger context (e.g. bge-large, 8192-token
#   models), raise this proportionally.
CHUNK_OVERLAP                  = int(os.getenv("CHUNK_OVERLAP", "50"))
# ↑ ~14% of CHUNK_SIZE. Just enough to avoid cutting a sentence at a boundary.
#   Do NOT set this high (e.g. 200) — with CHUNK_SIZE=350 that creates 57%
#   overlap which means near-duplicate chunks and a bloated index.
TOP_K                          = int(os.getenv("TOP_K", "10"))
# ↑ 10 is the sweet spot for 100+ docs. Higher = more noise sent to LLM.
CHUNKS_OUTPUT_FILE             = os.getenv("CHUNKS_OUTPUT_FILE", "chunks_output.txt")
SSL_VERIFY                     = os.getenv("SSL_VERIFY", "true").lower() == "true"
PARALLEL_WORKERS               = int(os.getenv("PARALLEL_WORKERS", "5"))
CHECKPOINT_EVERY               = int(os.getenv("CHECKPOINT_EVERY", "100"))
EMBEDDING_BATCH_SIZE           = int(os.getenv("EMBEDDING_BATCH_SIZE", "64"))
OLLAMA_MAX_RETRIES             = int(os.getenv("OLLAMA_MAX_RETRIES", "3"))
OLLAMA_RETRY_DELAY             = float(os.getenv("OLLAMA_RETRY_DELAY", "2.0"))
OLLAMA_NUM_PREDICT             = int(os.getenv("OLLAMA_NUM_PREDICT", "8192"))
# ↑ Max tokens the LLM can generate in one response.
#   Ollama default is 128 — way too small for multi-section answers.
#   4096 ensures full answers through all sections without cutoff.
DOC_MATCH_THRESHOLD            = float(os.getenv("DOC_MATCH_THRESHOLD", "0.90"))
# ↑ Legacy threshold — still used as a kill-switch (1.0 = disable filtering).
#   The actual matching now uses multiple methods with separate thresholds
#   (see MEANINGFUL_OVERLAP_THRESHOLD, FUZZY_THRESHOLD below).

# --- Document Name Matching — v2 Tunable Parameters ---
MEANINGFUL_OVERLAP_THRESHOLD   = float(os.getenv("MEANINGFUL_OVERLAP_THRESHOLD", "0.55"))
# ↑ Fraction of meaningful (non-code, non-filler) words from the document
#   name that must appear in the query for a KEYWORD match.
#   0.55 = 55% — e.g. "supplier risk assessment" has 3 meaningful words,
#   if 2 appear in the query (67%) that's a match.

FUZZY_THRESHOLD                = float(os.getenv("FUZZY_THRESHOLD", "0.70"))
# ↑ Minimum SequenceMatcher ratio for a FUZZY match on the clean document
#   name (after stripping code prefixes/suffixes) vs the clean query.
#   Only applies when the stem has 3+ meaningful words.

FUZZY_THRESHOLD_SHORT          = float(os.getenv("FUZZY_THRESHOLD_SHORT", "0.80"))
# ↑ Higher fuzzy threshold for short document names (≤2 meaningful words).
#   Prevents false positives like "collaboration tools policy" matching
#   "laptop policy" — the short name needs a closer match.

MAX_GENERIC_MATCHES            = int(os.getenv("MAX_GENERIC_MATCHES", "3"))
# ↑ If a query matches this many DIFFERENT documents (not versions of the
#   same doc), the filter stays OFF — the query is too generic.
#   e.g. "policy" would match 7+ docs -> filter OFF.

# --- Code Prefix / Suffix Regexes for Document Name Cleaning ---
# Strips corporate doc-code prefixes like: MG-CSS-GL-IT-001_, FS_PRP_08,
# HR-SOP-07, WI-05, SM-F-05, SOP-
_CODE_PREFIX_RE = re.compile(
    r"^"
    r"(?:"
    r"(?:[a-z]{1,4}[-_]){1,5}"   # e.g. mg-css-gl-it- or fs_prp_
    r"\d{1,4}"                    # e.g. 001, 08, 07
    r"[-_\s]+"                    # separator
    r")+"
)

# Strips trailing codes/dates like: _FS.SP 4.00_RA-01 28.02.23, v1
_CODE_SUFFIX_RE = re.compile(
    r"[_\s]+"
    r"(?:"
    r"fs\.sp[\s_]*\d+.*"         # _FS.SP 4.00_RA-01 ...
    r"|ra[\s_-]*\d+.*"           # _RA-01 28.02.23
    r"|v\d+\s*$"                  # v1 (at end)
    r"|\d{2}\.\d{2}(\.\d{2,4})?\s*$"  # 28.02.23 or 15.02 (at end)
    r")"
    , re.IGNORECASE
)

# Filler / stop words stripped from queries before matching
_FILLER_WORDS = frozenset({
    "what", "is", "the", "a", "an", "of", "in", "on", "for", "to", "from",
    "about", "tell", "me", "can", "you", "please", "how", "does", "do",
    "are", "was", "were", "be", "been", "being", "has", "have", "had",
    "that", "this", "these", "those", "which", "who", "whom", "whose",
    "with", "by", "at", "or", "and", "but", "not", "no", "if", "then",
    "there", "here", "it", "its", "i", "my", "we", "our", "they", "their",
    "show", "give", "list", "explain", "describe", "find", "usage", "use",
})

# Generic domain words that alone don't identify a specific document
_GENERIC_DOC_WORDS = frozenset({
    "policy", "sop", "procedure", "form", "guideline", "standard",
    "manual", "plan", "report", "matrix", "assessment",
})
BLOB_MAX_RETRIES               = int(os.getenv("BLOB_MAX_RETRIES", "3"))
FAISS_IVF_NLIST                = int(os.getenv("FAISS_IVF_NLIST", "100"))
FAISS_IVF_THRESHOLD            = int(os.getenv("FAISS_IVF_THRESHOLD", "5000"))
FAISS_NPROBE                   = int(os.getenv("FAISS_NPROBE", "16"))
DEBUG_CHUNKS_FILE              = os.getenv("DEBUG_CHUNKS_FILE", "debug_chunks_to_llm.txt")

# Minimum ratio of printable ASCII characters (0–1) — filters binary/corrupt .doc content
# Sections below this ratio are garbage (e.g., olefile binary extraction from bad .doc files)
MIN_ASCII_RATIO                = float(os.getenv("MIN_ASCII_RATIO", "0.50"))


# ---------------------------------------------------------------------------
# Chunk quality helpers — used at ingestion AND retrieval time
# ---------------------------------------------------------------------------
def _chunk_content_hash(text: str) -> str:
    """MD5 hash of stripped text — used for near-duplicate detection.

    For table rows (lines starting with |), we hash EACH row independently
    and combine them. This prevents a large table chunk from being treated
    as a duplicate just because most rows are the same — if even one row
    differs, the hash differs.
    """
    stripped = text.strip()
    # Check if this looks like a markdown table
    lines = stripped.split('\n')
    table_lines = [l for l in lines if l.strip().startswith('|')]
    if len(table_lines) >= 2:
        # It's a table — hash using source + row content to preserve row-level uniqueness
        # Include a length prefix so tables of different sizes get different hashes
        content = f"TABLE:{len(table_lines)}:" + ''.join(sorted(table_lines))
        return hashlib.md5(content.encode("utf-8", errors="ignore")).hexdigest()
    return hashlib.md5(stripped.encode("utf-8", errors="ignore")).hexdigest()


def _is_valid_section(text: str) -> bool:
    """
    Return True if a section/chunk is worth indexing.

    Only rejects binary garbage — sections with mostly non-ASCII characters
    that come from corrupt .doc piece-table extraction (e.g. raw OLE bytes
    decoded as unicode produce strings like 쌀쐀팀؀ etc.).

    Does NOT enforce a minimum length — even a short section like a form
    field header is valid content and should be kept.
    """
    clean = text.strip()
    if not clean:
        return False
    # Count printable ASCII chars (space through tilde) + common whitespace
    printable = sum(1 for c in clean if 32 <= ord(c) < 127 or c in '\t\n\r')
    if printable / len(clean) < MIN_ASCII_RATIO:
        return False
    return True


# ===========================================================================
# Azure Blob Helper — with retry logic
# ===========================================================================
class AzureBlobManager:
    """Streams documents from Azure Blob Storage — no local file downloads."""

    def __init__(self, connection_string: str, container_name: str):
        self.client = BlobServiceClient.from_connection_string(connection_string)
        self.container_client = self.client.get_container_client(container_name)
        self.container_name = container_name

    def list_blobs(self, prefix: Optional[str] = None) -> List[str]:
        return [b.name for b in self.container_client.list_blobs(name_starts_with=prefix)]

    def stream_blob(self, blob_name: str, max_retries: int = BLOB_MAX_RETRIES) -> bytes:
        """Download blob bytes with retry logic."""
        for attempt in range(1, max_retries + 1):
            try:
                blob_client = self.container_client.get_blob_client(blob_name)
                return blob_client.download_blob().readall()
            except Exception as e:
                if attempt == max_retries:
                    raise
                logger.warning(f"Blob download retry {attempt}/{max_retries} for {blob_name}: {e}")
                time.sleep(attempt * 1.5)

    def upload_bytes(self, blob_name: str, data: bytes):
        blob_client = self.container_client.get_blob_client(blob_name)
        blob_client.upload_blob(data, overwrite=True)

    def download_bytes(self, blob_name: str) -> bytes:
        blob_client = self.container_client.get_blob_client(blob_name)
        return blob_client.download_blob().readall()

    def delete_blob(self, blob_name: str):
        """Delete a blob — used by reset() to clear the stored index."""
        blob_client = self.container_client.get_blob_client(blob_name)
        if blob_client.exists():
            blob_client.delete_blob()

    def blob_exists(self, blob_name: str) -> bool:
        blob_client = self.container_client.get_blob_client(blob_name)
        return blob_client.exists()


# ===========================================================================
# Ingestion Manifest — tracks which blobs have been processed
# ===========================================================================
class IngestionManifest:
    """
    Tracks which documents have already been ingested.
    Enables incremental ingestion — skip already-processed blobs.
    """

    def __init__(self, manifest_path: str):
        self.manifest_path = manifest_path
        self._ingested: Dict[str, Dict[str, Any]] = {}
        self._load()

    def _load(self):
        if os.path.exists(self.manifest_path):
            try:
                with open(self.manifest_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                self._ingested = data.get("ingested", {})
                logger.info(f"Manifest loaded: {len(self._ingested)} previously ingested documents")
            except Exception as e:
                logger.warning(f"Could not load manifest: {e}")
                self._ingested = {}

    def save(self):
        with open(self.manifest_path, "w", encoding="utf-8") as f:
            json.dump({
                "ingested": self._ingested,
                "last_updated": datetime.now().isoformat(),
            }, f, indent=2)
        logger.info(f"Manifest saved: {len(self._ingested)} documents tracked")

    def is_ingested(self, blob_name: str) -> bool:
        return blob_name in self._ingested

    def mark_ingested(self, blob_name: str, chunk_count: int = 0, char_count: int = 0):
        self._ingested[blob_name] = {
            "ingested_at": datetime.now().isoformat(),
            "chunk_count": chunk_count,
            "char_count": char_count,
        }

    def get_unprocessed(self, blob_names: List[str]) -> List[str]:
        """Return only blob names that haven't been ingested yet."""
        return [b for b in blob_names if b not in self._ingested]

    @property
    def count(self) -> int:
        return len(self._ingested)


# ===========================================================================
# Chunk Output — Console + File
# ===========================================================================
class ChunkOutputWriter:
    """
    Writes chunk summaries to both the console and a chunks_output.txt file.
    Accumulates chunks across files and writes the final output on demand.
    """

    def __init__(self, output_path: str = CHUNKS_OUTPUT_FILE):
        self.output_path = output_path
        self._lines: List[str] = []
        self._total_chunks: int = 0

    def _format_chunks(self, chunks: List[Dict[str, Any]], verbose: bool = False) -> str:
        """Format chunk summaries into a string block."""
        if not chunks:
            return "\n  [No chunks generated]\n"

        role_counts = {}
        for c in chunks:
            role = c.get("chunk_role", "standalone")
            role_counts[role] = role_counts.get(role, 0) + 1

        lines = []
        lines.append(f"\n{'='*70}")
        lines.append(f"  CHUNK OUTPUT — {len(chunks)} chunk(s) generated")
        lines.append(f"  Breakdown: {', '.join(f'{v} {k}' for k, v in role_counts.items())}")
        lines.append(f"{'='*70}")

        for c in chunks:
            chunk_id   = c.get("chunk_id", "?")
            role       = c.get("chunk_role", "standalone")
            source     = c.get("source_file", "?")
            heading    = c.get("heading_text", "")
            parent_path = c.get("parent_path", "")
            hlevel     = c.get("heading_level", 0)
            parent_id  = c.get("parent_id")
            child_ids  = c.get("child_ids", [])
            text_len   = len(c.get("text", ""))
            text_preview = c.get("text", "")[:120].replace("\n", " ")

            role_icon = {"parent": "P", "child": "C", "standalone": "S"}.get(role, "?")

            lines.append(f"\n  Chunk [{chunk_id}] ({role_icon}) | Source: {source}")
            if heading:
                hierarchy = f"{parent_path} > {heading}" if parent_path else heading
                lines.append(f"    Section: {hierarchy} (H{hlevel})")
            if role == "parent" and child_ids:
                lines.append(f"    Children: {child_ids}")
            if role == "child" and parent_id is not None:
                lines.append(f"    Parent: Chunk [{parent_id}]")
            lines.append(f"    Length: {text_len} chars")

            if verbose:
                lines.append(f"    Text:\n    {c.get('text', '')}")
            else:
                lines.append(f"    Preview: {text_preview}{'...' if text_len > 120 else ''}")

        lines.append(f"\n{'='*70}\n")
        return "\n".join(lines)

    def write_chunks(self, chunks: List[Dict[str, Any]], verbose: bool = False):
        """Print chunks to console and accumulate for file output."""
        text = self._format_chunks(chunks, verbose)
        logger.info(text)
        self._lines.append(text)
        self._total_chunks += len(chunks)

    def save_to_file(self):
        """Write all accumulated chunk output to chunks_output.txt."""
        with open(self.output_path, "w", encoding="utf-8") as f:
            f.write(f"RAG Pipeline — Chunks Output\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Total Chunks: {self._total_chunks}\n")
            f.write(f"{'='*70}\n")
            for line in self._lines:
                f.write(line)
            f.write(f"\n\nTotal chunks in this run: {self._total_chunks}\n")
        logger.info(f"Chunks output saved to: {os.path.abspath(self.output_path)}")


# ===========================================================================
# Document Processor — with Hierarchical Parent-Child Chunking
# ===========================================================================
class DocumentProcessor:
    """
    Loads and chunks documents from bytes (Azure Blob) or local files.

    DOCX files use HIERARCHICAL parent-child chunking based on heading styles.
    DOC files use LibreOffice to convert to DOCX, then hierarchical parsing.
    PDF files use smart heading detection with font-size and heuristic analysis.
    ODT files use XML heading element detection.
    Other formats use sliding-window chunking.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf":  "pdf",
        ".doc":  "doc_libreoffice",
        ".docx": "docx",
        ".docm": "docx",
        ".dot":  "doc_libreoffice",
        ".dotx": "docx",
        ".odt":  "odt",
        ".rtf":  "rtf",
        ".xlsx": "xlsx",
        ".xlsm": "xlsx",
        ".xls":  "xls_legacy",
        ".csv":  "csv",
        ".txt":  "txt",
    }

    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4",
                      "Heading 5", "Heading 6", "heading 1", "heading 2",
                      "heading 3", "heading 4", "heading 5", "heading 6"}

    # -------------------------------------------------------------------
    # Load from bytes (Azure Blob streaming — no local file)
    # -------------------------------------------------------------------
    def load_from_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        doc_kind = self.SUPPORTED_EXTENSIONS.get(ext)
        if not doc_kind:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(self.SUPPORTED_EXTENSIONS.keys())}")

        dispatch = {
            "pdf":        self._load_pdf_bytes,
            "doc_libreoffice": self._load_doc_libreoffice,
            "docx":       self._load_docx_bytes,
            "odt":        self._load_odt_bytes,
            "rtf":        self._load_rtf_bytes,
            "xlsx":       self._load_xlsx_bytes,
            "xls_legacy": self._load_xls_bytes,
            "csv":        self._load_csv_bytes,
            "txt":        self._load_txt_bytes,
        }
        return dispatch[doc_kind](data, filename)

    # -------------------------------------------------------------------
    # Load from local file path
    # -------------------------------------------------------------------
    def load_document(self, file_path: str) -> List[Dict[str, Any]]:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        doc_kind = self.SUPPORTED_EXTENSIONS.get(ext)
        if not doc_kind:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(self.SUPPORTED_EXTENSIONS.keys())}")
        with open(file_path, "rb") as f:
            data = f.read()
        return self.load_from_bytes(data, os.path.basename(file_path))

    # ===================================================================
    # PDF — Docling Parse (primary) + pdfplumber fallback
    # ===================================================================

    def _load_pdf_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parse PDF with heading hierarchy awareness AND table extraction.

        Strategy cascade:
          1. Docling Parse — uses the same C++ PDF parser as IBM Docling.
             Extracts text lines with font metadata AND the PDF's embedded
             Table of Contents (bookmarks) for accurate heading detection.
          2. pdfplumber fallback — font-size clustering + heuristic rules
          3. Heuristic-only fallback — for scanned/text-only PDFs

        After text extraction, pdfplumber is used to extract tables with
        proper cell structure (works for PDFs with visible table grid lines).
        Tables are linked to their parent heading sections.

        Finally, page coverage is guaranteed — every page gets at least
        one section/chunk, even if heading detection missed it.
        """
        sections = []

        # --- Try Docling Parse first ---
        try:
            sections = self._load_pdf_docling(data, filename)
            if sections:
                heading_count = sum(1 for s in sections if s.get("heading_level", 0) > 0)
                logger.info(f"  PDF (Docling): {len(sections)} sections, {heading_count} with headings detected")
        except Exception as e:
            logger.warning(f"  PDF: Docling parse failed ({e}), falling back to pdfplumber")

        # --- Fallback: pdfplumber with font detection ---
        if not sections:
            try:
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    all_lines = self._extract_pdf_lines_with_fonts(pdf)
                    if all_lines:
                        sections = self._build_pdf_hierarchy(all_lines, filename)
                        if sections:
                            heading_count = sum(1 for s in sections if s.get("heading_level", 0) > 0)
                            logger.info(f"  PDF (pdfplumber): {len(sections)} sections, {heading_count} with headings detected")

                    if not sections:
                        logger.info("  PDF: Font-based detection yielded no headings, using heuristic fallback")
                        sections = self._load_pdf_fallback(pdf, filename)
            except Exception as e:
                logger.warning(f"  PDF: pdfplumber failed ({e}), using basic extraction")
                sections = self._load_pdf_basic(data, filename)

        # --- Extract tables from PDF (regardless of which text method worked) ---
        # pdfplumber's table extraction works on PDFs with visible grid lines.
        # This catches form-like and tabular content that text-only extraction misses.
        try:
            table_sections = self._extract_pdf_tables_raw(data, filename)
            if table_sections:
                # Remove text sections that are clearly table content (duplicate)
                # to avoid sending garbage like "CompanyAcme Corp" as a heading
                sections = self._deduplicate_pdf_table_content(sections, table_sections)
                # Now assign proper parent headings using cleaned sections
                for i, tbl in enumerate(table_sections):
                    tbl["parent_path"] = self._find_table_parent_heading(
                        sections, tbl.get("page_num", 1), i
                    )
                sections.extend(table_sections)
                logger.info(f"  PDF: {len(table_sections)} tables extracted")
        except Exception as e:
            logger.warning(f"  PDF: Table extraction failed ({e})")

        # --- PAGE COVERAGE GUARANTEE ---
        # Ensure every page in the PDF has at least one section/chunk.
        # If a page has NO sections (e.g., heading detection missed it,
        # or the page had no headings), extract its full text and add it
        # as a standalone section. This prevents the "no context" issue
        # when users ask about content on page 3/4 but those pages were
        # never chunked.
        try:
            sections = self._ensure_page_coverage(data, filename, sections)
        except Exception as e:
            logger.warning(f"  PDF: Page coverage check failed ({e})")

        return sections

    # ------------------------------------------------------------------
    # Docling Parse — Primary PDF Loader
    # ------------------------------------------------------------------
    def _load_pdf_docling(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Extract PDF content using Docling's C++ parser.

        Docling provides:
          - Text lines with font name, size, and confidence
          - The PDF's embedded Table of Contents (bookmarks/outline)
          - Accurate reading order
          - Cell-level text extraction

        We use the TOC for heading structure when available,
        and fall back to font-size clustering + heuristics when not.
        """
        parser = DoclingPdfParser()
        doc = parser.load(io.BytesIO(data), lazy=False)

        if not doc.is_loaded():
            return []

        num_pages = doc.number_of_pages()
        if num_pages == 0:
            doc.unload()
            return []

        # Step 1: Extract all text lines with font metadata
        all_lines = []
        for page_no in range(num_pages):
            try:
                page = doc.get_page(page_no)
            except Exception:
                continue

            if page is None:
                continue

            # Get textline-level cells (best balance of detail and grouping)
            for cell in page.textline_cells:
                text = cell.text.strip() if cell.text else ""
                if not text:
                    continue

                font_name = cell.font_name if cell.font_name else ""
                # Docling doesn't expose font_size directly on cells,
                # but we can derive it from the bounding box height
                bbox = cell.rect
                font_size = round(bbox.height, 1) if bbox and bbox.height > 0 else 12.0

                all_lines.append({
                    "text": text,
                    "font_size": font_size,
                    "font_name": font_name,
                    "page_num": page_no + 1,
                    "confidence": cell.confidence if cell.confidence else 1.0,
                })

        # Step 2: Try TOC-based heading detection
        toc = doc.get_table_of_contents()
        if toc and toc.children:
            sections = self._build_pdf_hierarchy_from_toc(all_lines, toc, filename)
            if sections:
                doc.unload()
                return sections

        # Step 3: Fallback to font-size + heuristic heading detection
        doc.unload()

        if not all_lines:
            return []

        return self._build_pdf_hierarchy(all_lines, filename)

    def _build_pdf_hierarchy_from_toc(
        self,
        lines: List[Dict[str, Any]],
        toc: Any,
        filename: str,
    ) -> List[Dict[str, Any]]:
        """
        Build hierarchical sections using the PDF's embedded TOC (bookmarks).

        The TOC provides accurate heading text and nesting levels directly
        from the document author — much more reliable than font-size guessing.

        Strategy:
          - Flatten the TOC into (title, level) pairs
          - For each TOC entry, find its content between this entry and the next
          - Build sections with proper parent-child paths
        """
        # Flatten TOC tree into ordered list of (title, level)
        toc_entries = []
        self._flatten_toc(toc, toc_entries)

        if not toc_entries:
            return []

        # For each TOC entry, collect text lines that belong to it
        # Since we don't have page positions from TOC, we use text matching:
        # find the TOC heading text in the lines and assign everything
        # until the next TOC heading
        sections = []
        heading_stack = []
        current_section_lines = []
        current_heading_level = 0
        current_heading_text = ""
        current_parent_path = ""

        # Build a set of all TOC heading texts for fast lookup
        toc_texts = set(entry[0].strip().lower() for entry in toc_entries)
        toc_level_map = {entry[0].strip().lower(): min(entry[1], 3) for entry in toc_entries}

        for line in lines:
            text = line["text"].strip()
            text_lower = text.lower()
            page_num = line.get("page_num", 1)

            # Check if this line matches a TOC heading
            matched_level = 0
            if text_lower in toc_level_map:
                matched_level = toc_level_map[text_lower]
            else:
                # Partial match: TOC entry might be a substring of the line
                for toc_text, toc_lvl in toc_level_map.items():
                    if toc_text and len(toc_text) > 5 and toc_text in text_lower:
                        matched_level = toc_lvl
                        break

            if matched_level > 0:
                # Flush previous section
                if current_section_lines:
                    section_text = "\n".join(current_section_lines)
                    sections.append({
                        "text": section_text, "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": current_heading_level,
                        "heading_text": current_heading_text,
                        "parent_path": current_parent_path,
                        "page_num": page_num,
                    })

                # Update heading stack
                while heading_stack and heading_stack[-1][0] >= matched_level:
                    heading_stack.pop()
                heading_stack.append((matched_level, text))

                current_heading_level = matched_level
                current_heading_text = text
                current_parent_path = " > ".join([h[1] for h in heading_stack[:-1]])
                current_section_lines = [text]
            else:
                current_section_lines.append(text)

        # Flush last section
        if current_section_lines:
            section_text = "\n".join(current_section_lines)
            sections.append({
                "text": section_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": current_heading_level,
                "heading_text": current_heading_text,
                "parent_path": current_parent_path,
            })

        # If no sections with headings, return empty (will trigger fallback)
        heading_count = sum(1 for s in sections if s.get("heading_level", 0) > 0)
        if heading_count == 0:
            return []

        return sections

    def _flatten_toc(self, toc_node: Any, result: List[Tuple[str, int]], level: int = 1):
        """Recursively flatten the TOC tree into (title, level) pairs."""
        if toc_node is None:
            return
        text = toc_node.text.strip() if hasattr(toc_node, 'text') and toc_node.text else ""
        if text:
            result.append((text, level))
        if hasattr(toc_node, 'children') and toc_node.children:
            for child in toc_node.children:
                self._flatten_toc(child, result, level + 1)

    # ------------------------------------------------------------------
    # pdfplumber — Fallback PDF Loader (kept for when Docling fails)
    # ------------------------------------------------------------------
    def _extract_pdf_lines_with_fonts(self, pdf) -> List[Dict[str, Any]]:
        """
        Extract text lines from PDF with font size information (pdfplumber).

        Uses pdfplumber's character-level data to determine the dominant
        font size for each text line.
        """
        all_lines = []

        for page_num, page in enumerate(pdf.pages):
            chars = page.chars
            if not chars:
                continue

            line_groups = {}
            for char in chars:
                top_key = round(char.get("top", 0), 0)
                matched_key = None
                for existing_key in line_groups:
                    if abs(existing_key - top_key) < 2:
                        matched_key = existing_key
                        break
                if matched_key is not None:
                    line_groups[matched_key].append(char)
                else:
                    line_groups[top_key] = [char]

            sorted_tops = sorted(line_groups.keys())

            for top in sorted_tops:
                line_chars = line_groups[top]
                if not line_chars:
                    continue
                line_chars.sort(key=lambda c: c.get("x0", 0))
                text = "".join(c.get("text", "") for c in line_chars).strip()
                if not text:
                    continue

                size_counts: Dict[float, int] = {}
                for c in line_chars:
                    size = round(c.get("size", 12), 1)
                    size_counts[size] = size_counts.get(size, 0) + len(c.get("text", ""))
                dominant_size = max(size_counts, key=size_counts.get) if size_counts else 12.0

                all_lines.append({
                    "text": text,
                    "font_size": dominant_size,
                    "page_num": page_num + 1,
                    "top": top,
                })

        return all_lines

    def _build_pdf_hierarchy(
        self, lines: List[Dict[str, Any]], filename: str
    ) -> List[Dict[str, Any]]:
        """
        Build hierarchical sections from PDF lines using font-size clustering
        and heuristic heading classification.
        """
        if not lines:
            return []

        font_sizes = [l["font_size"] for l in lines]
        median_size = float(np.median(font_sizes))

        unique_sizes = sorted(set(font_sizes), reverse=True)
        heading_sizes = [s for s in unique_sizes if s > median_size * 1.1]

        size_to_level: Dict[float, int] = {}
        for i, size in enumerate(heading_sizes[:3]):
            size_to_level[size] = i + 1

        lengths = [len(l["text"]) for l in lines if len(l["text"]) > 0]
        avg_len = sum(lengths) / len(lengths) if lengths else 100

        classified_lines = []
        for line in lines:
            text = line["text"]
            font_size = line["font_size"]
            page_num = line.get("page_num", 1)

            font_heading = size_to_level.get(font_size, 0)
            heuristic_heading = self._classify_heading(text, avg_len)
            heading_level = max(font_heading, heuristic_heading)
            heading_level = min(heading_level, 3)

            classified_lines.append({
                "text": text,
                "heading_level": heading_level,
                "page_num": page_num,
                "font_size": font_size,
            })

        sections = []
        heading_stack = []
        current_section_lines = []
        current_heading_level = 0
        current_heading_text = ""
        current_parent_path = ""

        for cl in classified_lines:
            hlevel = cl["heading_level"]
            ptext = cl["text"]
            page_num = cl["page_num"]

            if hlevel > 0:
                if current_section_lines:
                    section_text = "\n".join(current_section_lines)
                    sections.append({
                        "text": section_text, "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": current_heading_level,
                        "heading_text": current_heading_text,
                        "parent_path": current_parent_path,
                        "page_num": page_num,
                    })

                while heading_stack and heading_stack[-1][0] >= hlevel:
                    heading_stack.pop()
                heading_stack.append((hlevel, ptext))

                current_heading_level = hlevel
                current_heading_text = ptext
                current_parent_path = " > ".join([h[1] for h in heading_stack[:-1]])
                current_section_lines = [ptext]
            else:
                current_section_lines.append(ptext)

        if current_section_lines:
            section_text = "\n".join(current_section_lines)
            sections.append({
                "text": section_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": current_heading_level,
                "heading_text": current_heading_text,
                "parent_path": current_parent_path,
            })

        if not sections and classified_lines:
            full_text = "\n".join(cl["text"] for cl in classified_lines)
            sections.append({
                "text": full_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": 0, "heading_text": "", "parent_path": "",
            })

        # --- Merge sections that have the same heading_text and parent_path ---
        # PDFs with running page headers (e.g. "3." repeated on page 4) can
        # cause the parser to create two separate sections for the same logical
        # section — one for page 3 content (3.1–3.3) and another for page 4
        # content (3.4–3.6). Merge them so the LLM sees the full section.
        merged_sections = []
        for sec in sections:
            if (merged_sections
                    and sec.get("heading_level", 0) > 0
                    and sec.get("heading_text") == merged_sections[-1].get("heading_text")
                    and sec.get("parent_path") == merged_sections[-1].get("parent_path")):
                # Same heading — append text to the previous section
                merged_sections[-1]["text"] += "\n\n" + sec["text"]
            else:
                merged_sections.append(sec)

        return merged_sections

    def _load_pdf_fallback(
        self, pdf, filename: str
    ) -> List[Dict[str, Any]]:
        """Fallback PDF loader using pdfplumber heuristic detection."""
        all_paragraphs = []
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                page_paras = [p.strip() for p in text.split("\n") if p.strip()]
                all_paragraphs.extend(page_paras)
        if not all_paragraphs:
            return []
        return self._detect_headings_in_text(all_paragraphs, filename)

    def _load_pdf_basic(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """Last-resort PDF loader — just extract raw text."""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return [{"text": text, "source_file": filename, "doc_type": "narrative",
                         "heading_level": 0, "heading_text": "", "parent_path": ""}]
        except ImportError:
            pass
        return [{"text": "[Could not extract text from PDF]", "source_file": filename,
                 "doc_type": "narrative", "heading_level": 0, "heading_text": "", "parent_path": ""}]

    # ------------------------------------------------------------------
    # PDF Table Extraction — pdfplumber extract_tables()
    # ------------------------------------------------------------------
    def _extract_pdf_tables_raw(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF using pdfplumber's extract_tables().

        This catches form-like and tabular content that text-only extraction
        misses — tables with visible grid lines get proper cell structure.

        Returns list of section dicts with doc_type="tabular" and text
        formatted as markdown tables.
        """
        table_sections = []

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_num = page_idx + 1
                try:
                    tables = page.extract_tables()
                except Exception:
                    continue

                if not tables:
                    continue

                for table_idx, table in enumerate(tables):
                    if not table or len(table) < 2:
                        continue

                    # Convert to markdown table format
                    rows = []
                    for row in table:
                        cells = [
                            (cell or "").strip().replace("\n", " ")
                            for cell in row
                        ]
                        rows.append("| " + " | ".join(cells) + " |")

                    if len(rows) < 2:
                        continue

                    # Add separator after header row
                    header = rows[0]
                    num_cols = len(table[0])
                    separator = "| " + " | ".join(["---"] * num_cols) + " |"
                    body = rows[1:]

                    markdown_table = f"{header}\n{separator}\n" + "\n".join(body)

                    table_sections.append({
                        "text": markdown_table,
                        "source_file": filename,
                        "doc_type": "tabular",
                        "heading_level": 0,
                        "heading_text": f"Table {len(table_sections) + 1}",
                        "parent_path": "",
                        "page_num": page_num,
                        "table_idx": len(table_sections) + 1,
                    })

        return table_sections

    def _deduplicate_pdf_table_content(
        self,
        sections: List[Dict[str, Any]],
        table_sections: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        """
        Remove text sections whose content is clearly just a garbled table.

        When pdfplumber extracts text from a table page, the result is often
        garbage like "CompanyAcme CorpLocationNew York" — all cells run
        together with no spaces. The proper table extraction produces clean
        markdown. So we remove the garbled text versions.
        """
        # Collect all table texts (lowercase, stripped) for comparison
        table_texts = set()
        for tbl in table_sections:
            # Get just the cell content from the markdown table
            for line in tbl.get("text", "").split("\n"):
                if line.startswith("|"):
                    cells = [c.strip() for c in line.split("|") if c.strip()]
                    for cell in cells:
                        if len(cell) > 3:
                            table_texts.add(cell.lower())

        if not table_texts:
            return sections

        cleaned = []
        for sec in sections:
            # Skip sections that are primarily table content garbage
            text = sec.get("text", "")
            if sec.get("doc_type") == "tabular":
                cleaned.append(sec)
                continue

            # Check if this section's text is mostly table cell content
            text_words = text.lower().split()
            if len(text_words) < 5:
                cleaned.append(sec)
                continue

            # Count how many table cells appear in this text
            matched = sum(1 for tt in table_texts if tt in text.lower())
            # If most of the text is table cell content, it's a garbled duplicate
            if matched > len(table_texts) * 0.5 and len(text_words) < 100:
                logger.info(f"  PDF: Removed garbled table text section: '{text[:60]}...'")
                continue

            cleaned.append(sec)

        return cleaned

    def _find_table_parent_heading(
        self,
        sections: List[Dict[str, Any]],
        page_num: int,
        table_idx: int,
    ) -> str:
        """
        Find the closest heading to a table's page for parent context.

        Looks backward first (heading on or before the page), then forward
        if nothing found. Returns the heading hierarchy string.
        """
        heading_by_page = {}
        for sec in sections:
            hlevel = sec.get("heading_level", 0)
            if hlevel > 0 and sec.get("heading_text"):
                pn = sec.get("page_num", 0)
                if pn > 0:
                    heading = sec.get("heading_text", "")
                    parent = sec.get("parent_path", "")
                    hierarchy = f"{parent} > {heading}" if parent else heading
                    heading_by_page[pn] = hierarchy

        return self._find_nearest_heading(heading_by_page, page_num)

    # ------------------------------------------------------------------
    # Page Coverage Guarantee — no page left behind
    # ------------------------------------------------------------------
    def _ensure_page_coverage(
        self,
        data: bytes,
        filename: str,
        sections: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        """
        Guarantee that every page in the PDF has at least one section.

        Problem: The heading-based section builder can miss entire pages
        if no heading is detected on that page. For example:
          - Page 1: H1 detected -> section created
          - Page 2: Text under H1 -> included in page 1's section
          - Page 3: No heading, text didn't get grouped -> MISSING
          - Page 4: No heading, text didn't get grouped -> MISSING

        When a user asks about content on page 4, the LLM says "not
        enough context" because page 4 was never chunked.

        Fix: After all section building, check which page numbers are
        covered. For any uncovered page, extract its full text using
        pdfplumber and add it as a standalone section.
        """
        if not data:
            return sections

        # Step 1: Find which pages are covered by existing sections
        covered_pages = set()
        for sec in sections:
            page_num = sec.get("page_num")
            if page_num and page_num > 0:
                covered_pages.add(page_num)

        # Step 2: Get total page count from PDF
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                total_pages = len(pdf.pages)

                if total_pages == 0:
                    return sections

                # Step 3: Find uncovered pages
                all_pages = set(range(1, total_pages + 1))
                uncovered_pages = sorted(all_pages - covered_pages)

                if not uncovered_pages:
                    logger.info(f"  PDF: Page coverage OK — all {total_pages} pages covered")
                    return sections

                # Extra check: a page may be "covered" inside a hierarchical
                # section even if it wasn't assigned a page_num tag.
                # Only add fallback pages that genuinely have no content in
                # any existing section (check by page text substring matching).
                truly_uncovered = []
                all_section_text = " ".join(s.get("text", "") for s in sections).lower()

                for page_num in uncovered_pages:
                    page_idx = page_num - 1
                    if page_idx >= len(pdf.pages):
                        continue
                    page = pdf.pages[page_idx]
                    page_text = page.extract_text()
                    if not page_text or not page_text.strip():
                        continue  # empty page — skip entirely, don't add placeholder
                    # Check if a meaningful snippet of this page is already in sections
                    sample = page_text.strip()[:120].lower()
                    if sample and sample in all_section_text:
                        continue  # already covered in a section's text
                    truly_uncovered.append((page_num, page_text.strip()))

                if not truly_uncovered:
                    logger.info(f"  PDF: Page coverage OK — uncovered pages already in section text")
                    return sections

                logger.info(
                    f"  PDF: {len(truly_uncovered)} genuinely uncovered page(s) "
                    f"— extracting missing content"
                )

                # Step 4: Build heading map for assigning parent_path
                heading_by_page = {}
                for sec in sections:
                    hlevel = sec.get("heading_level", 0)
                    if hlevel > 0 and sec.get("heading_text"):
                        page_num = sec.get("page_num", 0)
                        if page_num > 0:
                            heading = sec.get("heading_text", "")
                            parent = sec.get("parent_path", "")
                            hierarchy = f"{parent} > {heading}" if parent else heading
                            heading_by_page[page_num] = hierarchy

                # Step 5: Add only genuinely uncovered pages (pre-filtered above)
                for page_num, page_text in truly_uncovered:
                    nearest_heading = self._find_nearest_heading(heading_by_page, page_num)
                    sections.append({
                        "text": page_text,
                        "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": 0,
                        "heading_text": "",
                        "parent_path": nearest_heading,
                        "page_num": page_num,
                    })

                logger.info(
                    f"  PDF: Page coverage fixed — added {len(truly_uncovered)} missing page(s)"
                )

        except Exception as e:
            logger.warning(f"  PDF: Page coverage check failed: {e}")

        return sections

    @staticmethod
    def _find_nearest_heading(
        heading_by_page: Dict[int, str],
        target_page: int,
    ) -> str:
        """
        Find the closest heading to a given page number.

        Looks backward first (heading on or before the page), then
        forward if nothing found. Returns the heading hierarchy string.
        """
        if not heading_by_page:
            return ""

        if target_page in heading_by_page:
            return heading_by_page[target_page]

        for page in range(target_page - 1, 0, -1):
            if page in heading_by_page:
                return heading_by_page[page]

        for page in range(target_page + 1, max(heading_by_page.keys()) + 1):
            if page in heading_by_page:
                return heading_by_page[page]

        return ""

    # ===================================================================
    # DOCX — Hierarchical Parent-Child Loading
    # ===================================================================

    def _load_docx_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX with heading hierarchy awareness.

        Tries Docling's MsWordBackend first — it uses outlineLvl + base_style
        for much better heading detection than python-docx. Falls back to
        python-docx if Docling is not available.
        """
        # --- Try Docling DOCX first ---
        if HAS_DOCLING_DOCX:
            try:
                sections = self._load_docx_docling(data, filename)
                if sections:
                    heading_count = sum(1 for s in sections if s.get("heading_level", 0) > 0)
                    logger.info(f"  DOCX (Docling): {len(sections)} sections, {heading_count} with headings detected")
                    return sections
            except Exception as e:
                logger.warning(f"  DOCX: Docling parse failed ({e}), falling back to python-docx")

        # --- Fallback: python-docx ---
        doc = DocxDocument(io.BytesIO(data))
        sections = self._parse_docx_hierarchy(doc, filename)

        # Also extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_docx_table(table)
            if table_text.strip():
                sections.append({
                    "text": table_text, "source_file": filename,
                    "table_idx": table_idx + 1, "doc_type": "tabular",
                    "heading_level": 0, "heading_text": f"Table {table_idx + 1}",
                    "parent_path": "",
                })

        return sections

    def _load_docx_docling(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parse DOCX using Docling's MsWordDocumentBackend.

        Advantages over python-docx:
          - Uses outlineLvl from the DOCX XML — catches headings that
            python-docx misses (e.g., custom-styled headings that are
            visually H1 but named "Title" or "Normal")
          - Uses base_style for heading classification
          - Better table extraction via export_to_markdown()
          - Detects DocItemLabel.SECTION_HEADER and TITLE properly

        Returns list of section dicts, or empty list if parsing fails.
        """
        # Write data to a temp file — Docling needs a file path
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        try:
            inp = InputDocument(
                path_or_stream=tmp_path,
                format=InputFormat.DOCX,
                filename=filename,
            )
            backend = MsWordDocumentBackend(inp)
            doc = backend.convert()

            if not doc:
                return []

            # Extract document items from Docling's document model
            sections = []
            for item, _ in doc.iterate_items():
                if not item.text or not item.text.strip():
                    continue

                # Determine heading level from Docling labels
                heading_level = 0
                heading_text = ""

                if item.label == DocItemLabel.TITLE:
                    heading_level = 1
                    heading_text = item.text.strip()
                elif item.label == DocItemLabel.SECTION_HEADER:
                    # Docling provides heading level through the item's
                    # position in the hierarchy. We infer it from the
                    # label — all section headers get level based on
                    # nesting, but Docling doesn't expose the exact level.
                    # We'll estimate from the text characteristics.
                    heading_level = 2  # Default to H2, could be refined
                    heading_text = item.text.strip()

                # Try to extract tables via Docling's markdown export
                # This produces much cleaner table output than python-docx
                text = item.text.strip()

                sections.append({
                    "text": text,
                    "source_file": filename,
                    "doc_type": "tabular" if item.label == DocItemLabel.TABLE else "narrative",
                    "heading_level": heading_level,
                    "heading_text": heading_text,
                    "parent_path": "",
                })

            # Try to get markdown export for better table formatting
            try:
                md_text = doc.export_to_markdown()
                if md_text and len(md_text) > sum(len(s["text"]) for s in sections):
                    # Markdown export is richer — use it to enhance table sections
                    # Find TABLE sections and try to replace with markdown version
                    for section in sections:
                        if section["doc_type"] == "tabular":
                            # Try to find this table in the markdown
                            table_text = section["text"]
                            if table_text[:30] in md_text:
                                # Find the markdown table block
                                start = md_text.find(table_text[:30])
                                if start >= 0:
                                    # Extract until the next blank line
                                    end = md_text.find("\n\n", start)
                                    if end == -1:
                                        end = len(md_text)
                                    md_table = md_text[start:end].strip()
                                    if len(md_table) > len(table_text):
                                        section["text"] = md_table
            except Exception:
                pass  # Non-critical — basic text extraction still works

            return sections

        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    def _parse_docx_hierarchy(self, doc: DocxDocument, filename: str) -> List[Dict[str, Any]]:
        """Walk through paragraphs, detect headings, group with parent-child paths."""
        para_info = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            heading_level = self._get_heading_level(style_name)
            para_info.append({
                "text": para.text.strip(),
                "heading_level": heading_level,
                "style_name": style_name,
            })

        sections = []
        heading_stack = []
        current_section_paras = []
        current_heading_level = 0
        current_heading_text = ""
        current_parent_path = ""

        for pinfo in para_info:
            hlevel = pinfo["heading_level"]
            ptext = pinfo["text"]

            if hlevel > 0:
                if current_section_paras:
                    section_text = "\n".join(current_section_paras)
                    sections.append({
                        "text": section_text, "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": current_heading_level,
                        "heading_text": current_heading_text,
                        "parent_path": current_parent_path,
                    })

                while heading_stack and heading_stack[-1][0] >= hlevel:
                    heading_stack.pop()
                heading_stack.append((hlevel, ptext))

                current_heading_level = hlevel
                current_heading_text = ptext
                current_parent_path = " > ".join([h[1] for h in heading_stack[:-1]])
                current_section_paras = [ptext]
            else:
                current_section_paras.append(ptext)

        if current_section_paras:
            section_text = "\n".join(current_section_paras)
            sections.append({
                "text": section_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": current_heading_level,
                "heading_text": current_heading_text,
                "parent_path": current_parent_path,
            })

        if not sections and para_info:
            for pinfo in para_info:
                sections.append({
                    "text": pinfo["text"], "source_file": filename,
                    "doc_type": "narrative",
                    "heading_level": 0, "heading_text": "", "parent_path": "",
                })

        return sections

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """Extract heading level from a Word style name. Returns 0 if not a heading."""
        if not style_name:
            return 0
        style_lower = style_name.lower().strip()
        if style_lower.startswith("heading"):
            parts = style_lower.split()
            if len(parts) == 2:
                try:
                    return int(parts[1])
                except ValueError:
                    pass
        if style_lower in ("title",):
            return 1
        if style_lower.startswith("heading") and style_lower[-1].isdigit():
            try:
                return int(style_lower[-1])
            except ValueError:
                pass
        return 0

    # ===================================================================
    # DOC — LibreOffice conversion to DOCX, then hierarchical parsing
    # ===================================================================

    def _load_doc_libreoffice(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Convert .doc to .docx using LibreOffice, then parse with the
        existing DOCX hierarchical parser.

        This gives MUCH better results than olefile because:
          - Heading styles are preserved (H1, H2, etc.)
          - Tables are properly extracted
          - No binary/garbage text from OLE parsing
          - Section hierarchy is accurate

        Falls back to olefile + mammoth if LibreOffice is not available.
        """
        # Try LibreOffice conversion first
        try:
            docx_data = self._convert_doc_to_docx(data, filename)
            if docx_data:
                # Parse the converted DOCX using the existing hierarchical parser
                result = self._load_docx_bytes(docx_data, filename)
                if result:
                    total_chars = sum(len(item["text"]) for item in result)
                    logger.info(f"  .doc -> .docx via LibreOffice ({total_chars} chars, {len(result)} sections)")
                    return result
        except Exception as e:
            logger.warning(f"  LibreOffice conversion failed for {filename}: {e}")

        # Fallback 1: olefile + piece-table
        logger.info(f"  Falling back to olefile for {filename}")
        try:
            result = self._doc_via_olefile_piecetable(data, filename)
            if result:
                total_chars = sum(len(item["text"]) for item in result)
                if total_chars > 10:
                    logger.info(f"  .doc loaded via olefile+piece-table ({total_chars} chars)")
                    return result
        except Exception as e:
            logger.warning(f"  olefile+piece-table failed: {e}")

        # Fallback 2: mammoth raw text extraction
        logger.info(f"  Falling back to mammoth for {filename}")
        try:
            result = self._doc_via_mammoth(data, filename)
            if result:
                total_chars = sum(len(item["text"]) for item in result)
                if total_chars > 10:
                    logger.info(f"  .doc loaded via mammoth ({total_chars} chars)")
                    return result
        except Exception as e:
            logger.warning(f"  mammoth fallback also failed: {e}")

        # Fallback 3: olefile raw stream
        try:
            result = self._doc_via_olefile_raw(data, filename)
            if result:
                total_chars = sum(len(item["text"]) for item in result)
                if total_chars > 10:
                    logger.info(f"  .doc loaded via olefile-raw ({total_chars} chars)")
                    return result
        except Exception as e:
            logger.warning(f"  olefile-raw fallback also failed: {e}")

        raise ValueError(f"Could not extract text from .doc file: {filename}")

    def _convert_doc_to_docx(self, data: bytes, filename: str) -> bytes:
        """
        Convert .doc bytes to .docx bytes using LibreOffice headless mode.

        Writes the .doc to a temp directory, runs LibreOffice to convert,
        then reads the resulting .docx file back as bytes.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            # Write the .doc file
            doc_path = os.path.join(tmpdir, filename)
            with open(doc_path, "wb") as f:
                f.write(data)

            # Run LibreOffice headless conversion
            result = subprocess.run(
                [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    doc_path,
                ],
                capture_output=True,
                timeout=60,
            )

            if result.returncode != 0:
                stderr = result.stderr.decode("utf-8", errors="ignore")
                raise RuntimeError(f"LibreOffice conversion failed (exit {result.returncode}): {stderr}")

            # Find the converted .docx file
            base_name = os.path.splitext(filename)[0]
            docx_path = os.path.join(tmpdir, base_name + ".docx")

            # Sometimes LibreOffice creates the file with a slightly different name
            if not os.path.exists(docx_path):
                # Search for any .docx in the temp dir
                docx_files = [f for f in os.listdir(tmpdir) if f.endswith(".docx")]
                if docx_files:
                    docx_path = os.path.join(tmpdir, docx_files[0])
                else:
                    raise RuntimeError(f"LibreOffice did not produce a .docx file for {filename}")

            # Read the converted .docx
            with open(docx_path, "rb") as f:
                return f.read()

    def _doc_via_olefile_piecetable(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parse .doc using olefile + Word Binary Format piece table.

        The Word .doc format stores text in a 'piece table' inside the
        WordDocument stream. We parse the FIB (File Information Block) to
        locate the piece table, extract clean text paragraphs, then apply
        smart heading detection heuristics.
        """
        ole = olefile.OleFileIO(io.BytesIO(data))

        if not ole.exists("WordDocument"):
            ole.close()
            raise ValueError("Not a valid .doc file — WordDocument stream missing")

        word_stream = ole.openstream("WordDocument").read()

        # Parse FIB to find the piece table
        paragraphs = self._extract_doc_paragraphs(ole, word_stream)
        ole.close()

        if not paragraphs:
            raise ValueError("No paragraphs extracted from piece table")

        # Apply smart heading detection
        sections = self._detect_headings_in_text(paragraphs, filename)
        return sections

    def _extract_doc_paragraphs(self, ole, word_stream: bytes) -> List[str]:
        """
        Extract paragraphs from a .doc file by parsing the Word Binary Format.

        Reads the FIB (File Information Block) to locate the piece table,
        then decodes the text pieces. Paragraph breaks are marked by
        ASCII 0x0D (carriage return) or 0x0D 0x0A sequences.
        """
        try:
            # FIB base fields
            wIdent = struct.unpack_from('<H', word_stream, 0)[0]
            if wIdent != 0xA5EC:
                raise ValueError(f"Invalid FIB magic: 0x{wIdent:04X}")

            nFib = struct.unpack_from('<H', word_stream, 2)[0]

            # Flags at offset 0x0A
            flags = struct.unpack_from('<H', word_stream, 0x0A)[0]
            fComplex = (flags >> 2) & 1  # bit 2

            # ccpText — character count of main document text
            # Offset varies by FIB version; for Word 97+, it's at 0x4C
            if len(word_stream) > 0x50:
                ccpText = struct.unpack_from('<i', word_stream, 0x4C)[0]
            else:
                ccpText = 0

            # Try to find the piece table through the Table stream
            # The clx structure in the Table stream contains the piece table
            piece_table_text = None

            for table_name in ["1Table", "0Table"]:
                if not ole.exists(table_name):
                    continue
                try:
                    table_stream = ole.openstream(table_name).read()
                    # Search for the clx structure (Pcdt signature = 0x02)
                    # The clx starts with optional Prc entries, then Pcdt
                    pos = 0
                    while pos < len(table_stream) - 4:
                        clxt = table_stream[pos]
                        if clxt == 0x01:
                            # Prc entry — skip it
                            cbGrpprl = struct.unpack_from('<H', table_stream, pos + 1)[0]
                            pos += 3 + cbGrpprl
                        elif clxt == 0x02:
                            # Pcdt found — this is the piece table
                            lcb = struct.unpack_from('<I', table_stream, pos + 1)[0]
                            pcd_data = table_stream[pos + 5: pos + 5 + lcb]
                            piece_table_text = self._decode_piece_table(pcd_data, word_stream)
                            break
                        else:
                            pos += 1
                    if piece_table_text is not None:
                        break
                except Exception:
                    continue

            if piece_table_text is None:
                # Fallback: try extracting text directly from WordDocument stream
                # using the character positions from FIB
                if ccpText > 0 and len(word_stream) > 0x100:
                    # Simple extraction: try to decode the document text area
                    # For non-complex documents, text follows the FIB
                    raise ValueError("Could not locate piece table")
                raise ValueError("No text extracted")

            # Split into paragraphs
            # Word uses \r (0x0D) as paragraph mark; \r\n = hard line break
            raw_text = piece_table_text
            # Normalize line endings
            raw_text = raw_text.replace('\r\n', '\n').replace('\r', '\n')
            paragraphs = [p.strip() for p in raw_text.split('\n') if p.strip()]

            return paragraphs

        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Piece table parsing failed: {e}")

    @staticmethod
    def _decode_piece_table(pcd_data: bytes, word_stream: bytes) -> str:
        """
        Decode text from the piece table (Pcd array).

        Each piece descriptor (Pcd) is 8 bytes:
          - 2 bytes: fc (file character position) + flag
          - 4 bytes: prm (property modifiers)
        The high bit of fc indicates text encoding:
          - 0 = uncompressed (UTF-16LE in WordDocument stream)
          - 1 = compressed (ANSI/CP1252, offset = fc/2)
        """
        n = len(pcd_data) // 8  # number of pieces
        text_parts = []

        for i in range(n):
            pcd_offset = i * 8
            fc_value = struct.unpack_from('<I', pcd_data, pcd_offset)[0]

            is_compressed = (fc_value & 0x40000000) != 0

            if is_compressed:
                # Compressed: ANSI text, offset is fc/2 in WordDocument stream
                byte_offset = (fc_value & 0x3FFFFFFF) // 2
                # Read a chunk of bytes (up to next piece or reasonable size)
                chunk_size = min(4096, len(word_stream) - byte_offset)
                if byte_offset < len(word_stream) and chunk_size > 0:
                    raw = word_stream[byte_offset:byte_offset + chunk_size]
                    # Decode as CP1252 (Windows Latin-1)
                    try:
                        text = raw.decode('cp1252', errors='ignore')
                    except Exception:
                        text = raw.decode('latin-1', errors='ignore')
                    # Clean up: remove non-printable chars except common whitespace
                    text = re.sub(r'[\x00-\x06\x08\x0e-\x1f]', '', text)
                    text_parts.append(text)
            else:
                # Uncompressed: UTF-16LE text in WordDocument stream
                byte_offset = fc_value & 0x3FFFFFFF
                chunk_size = min(8192, len(word_stream) - byte_offset)
                if byte_offset < len(word_stream) and chunk_size > 0:
                    raw = word_stream[byte_offset:byte_offset + chunk_size]
                    try:
                        text = raw.decode('utf-16-le', errors='ignore')
                    except Exception:
                        text = raw.decode('latin-1', errors='ignore')
                    text = re.sub(r'[\x00-\x06\x08\x0e-\x1f]', '', text)
                    text_parts.append(text)

        return "".join(text_parts)

    # ===================================================================
    # Shared heading detection (used by both .doc and .pdf)
    # ===================================================================

    def _detect_headings_in_doc_text(
        self, paragraphs: List[str], filename: str
    ) -> List[Dict[str, Any]]:
        """
        Smart heading detection for .doc text that has no style information.
        Alias for _detect_headings_in_text for backward compatibility.
        """
        return self._detect_headings_in_text(paragraphs, filename)

    def _detect_headings_in_text(
        self, paragraphs: List[str], filename: str
    ) -> List[Dict[str, Any]]:
        """
        Smart heading detection for plain text that has no style information.

        Used by both .doc and .pdf (heuristic fallback) loaders.

        Heuristics used:
          1. Numbered sections: "1.", "1.1", "2.3.1" etc.
          2. ALL CAPS lines (common in corporate documents)
          3. Short lines (< 80 chars) that look like titles
          4. Lines ending with a colon (often sub-headings)
          5. Lines that are significantly shorter than average paragraph length

        Detected headings are assigned heading levels and grouped into
        sections just like DOCX hierarchical parsing.
        """
        if not paragraphs:
            return []

        # --- Phase 1: Detect which lines are headings ---
        heading_info = []  # [(paragraph_text, heading_level), ...]

        # Calculate average paragraph length for heuristic
        lengths = [len(p) for p in paragraphs if len(p) > 0]
        avg_len = sum(lengths) / len(lengths) if lengths else 100

        for para in paragraphs:
            hlevel = self._classify_heading(para, avg_len)
            heading_info.append((para, hlevel))

        # --- Phase 2: Build sections with hierarchy (same as DOCX) ---
        sections = []
        heading_stack = []
        current_section_paras = []
        current_heading_level = 0
        current_heading_text = ""
        current_parent_path = ""

        for para, hlevel in heading_info:
            if hlevel > 0:
                # Flush previous section
                if current_section_paras:
                    section_text = "\n".join(current_section_paras)
                    sections.append({
                        "text": section_text, "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": current_heading_level,
                        "heading_text": current_heading_text,
                        "parent_path": current_parent_path,
                    })

                # Update heading stack
                while heading_stack and heading_stack[-1][0] >= hlevel:
                    heading_stack.pop()
                heading_stack.append((hlevel, para))

                current_heading_level = hlevel
                current_heading_text = para
                current_parent_path = " > ".join([h[1] for h in heading_stack[:-1]])
                current_section_paras = [para]
            else:
                current_section_paras.append(para)

        # Flush last section
        if current_section_paras:
            section_text = "\n".join(current_section_paras)
            sections.append({
                "text": section_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": current_heading_level,
                "heading_text": current_heading_text,
                "parent_path": current_parent_path,
            })

        # No headings detected — return flat paragraphs
        if not sections and paragraphs:
            full_text = "\n".join(paragraphs)
            sections.append({
                "text": full_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": 0, "heading_text": "", "parent_path": "",
            })

        # Merge sections with the same heading (cross-page splits)
        merged_sections = []
        for sec in sections:
            if (merged_sections
                    and sec.get("heading_level", 0) > 0
                    and sec.get("heading_text") == merged_sections[-1].get("heading_text")
                    and sec.get("parent_path") == merged_sections[-1].get("parent_path")):
                merged_sections[-1]["text"] += "\n\n" + sec["text"]
            else:
                merged_sections.append(sec)

        return merged_sections

    @staticmethod
    def _classify_heading(text: str, avg_para_len: float) -> int:
        """
        Classify a paragraph as a heading (1-3) or body text (0).

        Returns heading level or 0 for body text.

        Used by both .doc and .pdf heuristic heading detection.
        """
        stripped = text.strip()
        if not stripped:
            return 0

        # GUARD: bare number-only lines like "3." or "3" or "4." are running
        # page headers in PDFs — NOT section headings. Treat as body text.
        # Without this, "3." on page 4 creates a new spurious H1 section that
        # swallows 3.4–3.6 content into a misnamed parent chunk.
        if re.match(r'^\d+\.?\s*$', stripped):
            return 0

        # GUARD: document reference codes like "MG-CSS-GL-IT-001" that appear
        # as running headers on every page — not real section headings.
        if re.match(r'^[A-Z]{2,}-[A-Z0-9]{2,}-[A-Z]{2,}-[A-Z]{2,}-\d+', stripped):
            return 0

        # Heuristic 1: Numbered sections — "1. Title", "1.1 Title", etc.
        # Requires at least one word of title text after the number.
        numbered_match = re.match(r'^(\d+(?:\.\d+)*)\s+\S', stripped)
        if numbered_match:
            num_parts = numbered_match.group(1).split('.')
            level = len(num_parts)
            return min(level, 3)

        # Heuristic 2: Roman numeral sections — "I.", "II.", "III."
        roman_match = re.match(r'^(X{0,3}IX|X{0,3}IV|X{0,3}V?I{0,3})\.\s+\S', stripped, re.IGNORECASE)
        if roman_match:
            return 1

        # Heuristic 3: Lettered sections — "A.", "B.", "a)", "b)"
        letter_match = re.match(r'^([A-Z])[.\)]\s+\S', stripped)
        if letter_match:
            return 2

        # Heuristic 4: ALL CAPS lines (common in corporate docs)
        if stripped.isupper() and len(stripped) > 3 and len(stripped) < 120:
            # Shorter ALL CAPS = higher level heading
            if len(stripped) < 40:
                return 1
            return 2

        # Heuristic 5: Title Case lines that are short (proper nouns capitalized)
        words = stripped.split()
        if len(words) <= 10 and len(stripped) < 80:
            title_case_count = sum(1 for w in words if w[0].isupper())
            if title_case_count >= len(words) * 0.7 and len(words) >= 2:
                # Most words are capitalized — likely a heading
                if len(stripped) < 40:
                    return 2
                return 3

        # Heuristic 6: Short lines much shorter than average
        # Tightened: require at least 2 words to avoid single-word page headers
        if (len(stripped) < 60 and avg_para_len > 100
                and len(words) >= 2 and len(words) <= 8):
            return 3

        # Heuristic 7: Lines ending with a colon (sub-headings)
        if stripped.endswith(':') and len(stripped) < 80:
            return 2

        return 0

    def _doc_via_mammoth(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Mammoth .doc extraction with table support.

        ``mammoth.extract_raw_text`` silently drops all tables — it only walks
        paragraph nodes and ignores <w:tbl> elements entirely.

        We switch to ``mammoth.convert_to_html`` instead, which preserves
        tables as <table> HTML.  We then:
          1. Parse the HTML with the stdlib ``html.parser``
          2. Convert <table> elements to pipe-delimited markdown chunks
          3. Extract the remaining paragraph text (outside tables) as before

        Falls back to extract_raw_text if the HTML route fails.
        """
        import html as _html
        from html.parser import HTMLParser

        class _DocTableHTMLParser(HTMLParser):
            """Minimal HTML parser that extracts paragraphs and tables."""

            def __init__(self):
                super().__init__()
                self.sections: List[Dict] = []
                self._in_table = False
                self._in_cell = False
                self._current_row: List[str] = []
                self._current_cell_buf: List[str] = []
                self._current_table_rows: List[List[str]] = []
                self._para_buf: List[str] = []

            def handle_starttag(self, tag, attrs):
                if tag == "table":
                    self._flush_para()
                    self._in_table = True
                    self._current_table_rows = []
                elif tag == "tr":
                    self._current_row = []
                elif tag in ("td", "th"):
                    self._in_cell = True
                    self._current_cell_buf = []

            def handle_endtag(self, tag):
                if tag in ("td", "th"):
                    self._in_cell = False
                    self._current_row.append(
                        _html.unescape("".join(self._current_cell_buf)).strip()
                    )
                elif tag == "tr":
                    if self._current_row:
                        self._current_table_rows.append(self._current_row)
                elif tag == "table":
                    self._in_table = False
                    self._flush_table()
                elif tag in ("p", "br") and not self._in_table:
                    text = _html.unescape("".join(self._para_buf)).strip()
                    if text:
                        self.sections.append({"text": text, "is_table": False})
                    self._para_buf = []

            def handle_data(self, data):
                if self._in_cell:
                    self._current_cell_buf.append(data)
                elif not self._in_table:
                    self._para_buf.append(data)

            def _flush_para(self):
                text = _html.unescape("".join(self._para_buf)).strip()
                if text:
                    self.sections.append({"text": text, "is_table": False})
                self._para_buf = []

            def _flush_table(self):
                rows = self._current_table_rows
                if not rows:
                    return
                md_rows = ["| " + " | ".join(row) + " |" for row in rows]
                num_cols = max(len(r) for r in rows)
                separator = "| " + " | ".join(["---"] * num_cols) + " |"
                table_md = (
                    md_rows[0] + "\n" + separator + "\n" + "\n".join(md_rows[1:])
                    if len(md_rows) >= 1 else "\n".join(md_rows)
                )
                self.sections.append({"text": table_md, "is_table": True})

        try:
            html_result = mammoth.convert_to_html(io.BytesIO(data))
            html_text = html_result.value or ""

            if html_text.strip():
                parser = _DocTableHTMLParser()
                parser.feed(html_text)
                parser._flush_para()  # flush any trailing paragraph text

                raw_sections = parser.sections
                if raw_sections:
                    # Run heading detection across narrative paragraphs only,
                    # then splice table sections back in their original order.
                    narrative_texts = [
                        s["text"] for s in raw_sections if not s["is_table"]
                    ]
                    narrative_sections = (
                        self._detect_headings_in_text(narrative_texts, filename)
                        if narrative_texts else []
                    )

                    final: List[Dict[str, Any]] = []
                    narr_iter = iter(narrative_sections)
                    for sec in raw_sections:
                        if sec["is_table"]:
                            final.append({
                                "text": sec["text"],
                                "source_file": filename,
                                "doc_type": "tabular",
                                "heading_level": 0,
                                "heading_text": "",
                                "parent_path": "",
                            })
                        else:
                            ns = next(narr_iter, None)
                            if ns:
                                final.append(ns)
                    return final

        except Exception as e:
            logger.warning(
                f"  mammoth HTML conversion failed ({e}), falling back to extract_raw_text"
            )

        # Fallback: original raw-text path (tables will be missing, but safe)
        result = mammoth.extract_raw_text(io.BytesIO(data))
        text = result.value.strip()
        if not text:
            return []
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        return self._detect_headings_in_text(paragraphs, filename)

    def _doc_via_olefile_raw(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Last-resort .doc extraction using raw olefile stream parsing.

        The original implementation read the Table streams (1Table / 0Table)
        and decoded them as UTF-16-LE text, which produced garbled binary
        output and missed all table content entirely.

        This revision instead:
          1. Reads the ``WordDocument`` stream and attempts to parse it as
             UTF-16-LE XML so that Word Binary Format XML-based documents
             (Word 2003 XML / WordprocessingML) can yield table content.
          2. Falls back to scanning the raw bytes for printable text runs
             (the classic "strings"-style extraction) so we always get
             *something* even from fully binary Word 97-2003 .doc files.
          3. Separately attempts a lightweight XML parse of any embedded
             ``word/document.xml``-style content when the OLE container
             happens to wrap an OOXML document (rare, but possible).

        Tables found via XML are converted to pipe-delimited markdown and
        returned as ``doc_type="tabular"`` sections alongside the narrative
        paragraphs so the chunker and retriever see the full document.
        """
        import xml.etree.ElementTree as ET

        # ----------------------------------------------------------------
        # Helper: extract paragraphs and tables from WordprocessingML XML
        # ----------------------------------------------------------------
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def _xml_to_sections(xml_bytes: bytes) -> List[Dict[str, Any]]:
            """Parse WordprocessingML XML; return list of {text, is_table}."""
            try:
                root = ET.fromstring(xml_bytes)
            except ET.ParseError:
                return []

            results = []

            # Walk direct children of <w:body> (or the root if body absent)
            body = root.find(f"{{{W}}}body") or root
            for child in body:
                local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if local == "p":
                    # Paragraph — collect all <w:t> text runs
                    texts = [t.text or "" for t in child.iter(f"{{{W}}}t")]
                    text = "".join(texts).strip()
                    if text:
                        results.append({"text": text, "is_table": False})

                elif local == "tbl":
                    # Table — iterate rows and cells
                    rows_data: List[List[str]] = []
                    for tr in child.iter(f"{{{W}}}tr"):
                        row_cells: List[str] = []
                        seen_tc: set = set()
                        for tc in tr.iter(f"{{{W}}}tc"):
                            tc_id = id(tc)
                            if tc_id in seen_tc:
                                row_cells.append("")
                                continue
                            seen_tc.add(tc_id)
                            cell_texts = [
                                t.text or "" for t in tc.iter(f"{{{W}}}t")
                            ]
                            row_cells.append("".join(cell_texts).strip())
                        if row_cells:
                            rows_data.append(row_cells)

                    if rows_data:
                        md_rows = ["| " + " | ".join(r) + " |" for r in rows_data]
                        num_cols = max(len(r) for r in rows_data)
                        sep = "| " + " | ".join(["---"] * num_cols) + " |"
                        table_md = md_rows[0] + "\n" + sep + "\n" + "\n".join(md_rows[1:])
                        results.append({"text": table_md, "is_table": True})

            return results

        # ----------------------------------------------------------------
        # Helper: printable-text fallback (strings-style byte scan)
        # ----------------------------------------------------------------
        def _strings_extract(raw: bytes, min_run: int = 4) -> List[str]:
            """Extract printable ASCII runs of length >= min_run."""
            runs: List[str] = []
            buf: List[str] = []
            for b in raw:
                if 32 <= b < 127:
                    buf.append(chr(b))
                else:
                    if len(buf) >= min_run:
                        runs.append("".join(buf).strip())
                    buf = []
            if len(buf) >= min_run:
                runs.append("".join(buf).strip())
            return [r for r in runs if r]

        # ----------------------------------------------------------------
        # Main extraction logic
        # ----------------------------------------------------------------
        ole = olefile.OleFileIO(io.BytesIO(data))
        raw_sections: List[Dict] = []

        try:
            # Attempt 1 — parse WordDocument stream as XML
            if ole.exists("WordDocument"):
                wd_bytes = ole.openstream("WordDocument").read()
                xml_secs = _xml_to_sections(wd_bytes)
                if xml_secs:
                    raw_sections.extend(xml_secs)

            # Attempt 2 — scan Table streams for readable UTF-16-LE text
            # (covers binary Word 97-2003 .doc where WordDocument is binary)
            if not raw_sections:
                for stream_name in ["1Table", "0Table"]:
                    if ole.exists(stream_name):
                        try:
                            stream_data = ole.openstream(stream_name).read()
                            decoded = stream_data.decode("utf-16-le", errors="ignore")
                            cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", decoded)
                            paras = [p.strip() for p in cleaned.split("\n") if p.strip()]
                            for p in paras:
                                raw_sections.append({"text": p, "is_table": False})
                        except Exception:
                            continue

            # Attempt 3 — strings-style byte scan of all streams as last resort
            if not raw_sections and ole.exists("WordDocument"):
                wd_bytes = ole.openstream("WordDocument").read()
                strings = _strings_extract(wd_bytes)
                for s in strings:
                    raw_sections.append({"text": s, "is_table": False})

        finally:
            ole.close()

        if not raw_sections:
            return []

        # Build final section list: table sections become tabular chunks,
        # narrative paragraphs go through heading detection.
        table_secs: List[Dict[str, Any]] = []
        narrative_texts: List[str] = []

        for sec in raw_sections:
            if sec["is_table"]:
                table_secs.append({
                    "text": sec["text"],
                    "source_file": filename,
                    "doc_type": "tabular",
                    "heading_level": 0,
                    "heading_text": "",
                    "parent_path": "",
                })
            else:
                narrative_texts.append(sec["text"])

        narrative_secs = (
            self._detect_headings_in_text(narrative_texts, filename)
            if narrative_texts else []
        )

        # Interleave in original order (tables first within their position,
        # then narrative — a simple concat is fine here since olefile is
        # last-resort and positional fidelity is best-effort).
        return narrative_secs + table_secs

    # ===================================================================
    # ODT — with XML heading detection
    # ===================================================================

    def _load_odt_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Parse ODT with heading hierarchy awareness from XML elements.

        ODT files use <text:h text:outline-level="N"> elements for headings,
        which we detect and use for parent-child chunking.
        """
        import zipfile
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            with zf.open("content.xml") as content_file:
                tree = ET.parse(content_file)

        root = tree.getroot()
        ns = {
            "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
            "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
        }

        # Extract all text elements with heading level info
        elements = []
        for elem in root.iter():
            tag = elem.tag
            # Check if it's a heading element
            if tag.endswith("}h") or tag == "h":
                level_str = elem.get(
                    "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}outline-level",
                    "1"
                )
                try:
                    level = int(level_str)
                except (ValueError, TypeError):
                    level = 1
                text = "".join(elem.itertext()).strip()
                if text:
                    elements.append({"text": text, "heading_level": min(level, 3)})
            elif tag.endswith("}p") or tag == "p":
                text = "".join(elem.itertext()).strip()
                if text:
                    elements.append({"text": text, "heading_level": 0})

        if not elements:
            # Fallback: just grab all text:p elements without hierarchy
            paragraphs = root.findall(".//text:p", ns)
            content = []
            for para in paragraphs:
                if para.text and para.text.strip():
                    content.append({
                        "text": para.text.strip(), "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": 0, "heading_text": "", "parent_path": "",
                    })
            return content

        # Build hierarchy
        sections = []
        heading_stack = []
        current_section_paras = []
        current_heading_level = 0
        current_heading_text = ""
        current_parent_path = ""

        for elem in elements:
            hlevel = elem["heading_level"]
            ptext = elem["text"]

            if hlevel > 0:
                if current_section_paras:
                    section_text = "\n".join(current_section_paras)
                    sections.append({
                        "text": section_text, "source_file": filename,
                        "doc_type": "narrative",
                        "heading_level": current_heading_level,
                        "heading_text": current_heading_text,
                        "parent_path": current_parent_path,
                    })

                while heading_stack and heading_stack[-1][0] >= hlevel:
                    heading_stack.pop()
                heading_stack.append((hlevel, ptext))

                current_heading_level = hlevel
                current_heading_text = ptext
                current_parent_path = " > ".join([h[1] for h in heading_stack[:-1]])
                current_section_paras = [ptext]
            else:
                current_section_paras.append(ptext)

        if current_section_paras:
            section_text = "\n".join(current_section_paras)
            sections.append({
                "text": section_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": current_heading_level,
                "heading_text": current_heading_text,
                "parent_path": current_parent_path,
            })

        if not sections:
            full_text = "\n".join(e["text"] for e in elements)
            sections.append({
                "text": full_text, "source_file": filename,
                "doc_type": "narrative",
                "heading_level": 0, "heading_text": "", "parent_path": "",
            })

        return sections

    # ===================================================================
    # Other format loaders (non-hierarchical)
    # ===================================================================

    def _load_txt_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        text = data.decode("utf-8", errors="ignore")
        return [{"text": text, "source_file": filename, "doc_type": "narrative",
                 "heading_level": 0, "heading_text": "", "parent_path": ""}]

    def _load_csv_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        df = pd.read_csv(io.BytesIO(data))
        content = df.to_string(index=False)
        if not content.strip():
            return []
        return [{"text": content, "source_file": filename, "doc_type": "tabular",
                 "heading_level": 0, "heading_text": "", "parent_path": ""}]

    def _load_rtf_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        rtf_text = data.decode("utf-8", errors="ignore")
        text = re.sub(r"\\[a-z]+\d*\s?", "", rtf_text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\\", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            return []
        return [{"text": text, "source_file": filename, "doc_type": "narrative",
                 "heading_level": 0, "heading_text": "", "parent_path": ""}]

    def _load_xlsx_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Smart Excel parser for .xlsx files.

        Handles the messy real-world Excel files in the SAMCO corpus:
          - Merged cells (forward-filled via openpyxl)
          - Auto-detects header row (scans rows 0-15)
          - Filters empty/sparse sheets
          - Detects section separator rows (e.g. "Corrugation - Related")
          - Outputs markdown tables compatible with the existing chunker
        """
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        content = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Step 1: Read cells into a grid with merged cells forward-filled
            grid = self._xlsx_read_sheet_with_merges(ws)
            if not grid:
                continue

            # Step 2: Count non-null cells — skip very sparse sheets
            # Find columns that have at least one non-empty cell so that
            # empty formatting columns (common in HACCP template sheets
            # with 197+ columns but only 14 used) don't inflate sparsity.
            active_cols = set()
            filled_cells = 0
            for row in grid:
                for ci, cell in enumerate(row):
                    if cell is not None and str(cell).strip():
                        active_cols.add(ci)
                        filled_cells += 1

            if filled_cells < 8:
                logger.info(f"  Excel: skipping sparse sheet '{sheet_name}' "
                            f"({filled_cells} filled cells)")
                continue

            active_total = len(grid) * len(active_cols) if active_cols else 0
            if active_total > 0 and filled_cells < 15 and filled_cells / active_total < 0.10:
                logger.info(f"  Excel: skipping sparse sheet '{sheet_name}' "
                            f"({filled_cells}/{active_total} active cells filled)")
                continue

            # Step 3: Auto-detect header row
            header_idx = self._xlsx_detect_header_row(grid)
            if header_idx is None:
                logger.info(f"  Excel: no header detected in sheet '{sheet_name}', skipping")
                continue

            # Step 3b: Capture pre-header context (title, definitions, etc.)
            pre_header_lines = []
            for row in grid[:header_idx]:
                texts = [
                    str(c).strip() for c in row
                    if c is not None and str(c).strip()
                ]
                unique_texts = list(dict.fromkeys(texts))
                for t in unique_texts:
                    if len(t) > 1 and t not in pre_header_lines:
                        pre_header_lines.append(t)

            headers = [
                str(c).strip() if c is not None else ""
                for c in grid[header_idx]
            ]
            data_rows = grid[header_idx + 1:]

            # Step 4: Keep only columns that have a non-empty header
            non_empty_cols = [i for i, h in enumerate(headers) if h]
            if not non_empty_cols:
                continue
            headers = [headers[i] for i in non_empty_cols]
            data_rows = [
                [row[i] if i < len(row) else "" for i in non_empty_cols]
                for row in data_rows
            ]

            # Step 5: Drop rows that are completely empty
            data_rows = [
                row for row in data_rows
                if any(
                    c is not None and str(c).strip()
                    for c in row
                )
            ]
            if not data_rows:
                continue

            # Step 6: Detect section separators and split
            sections = self._xlsx_split_sections(headers, data_rows)

            # Step 6b: Empty sections are footer metadata — merge into context
            footer_lines = []
            live_sections = []
            for section_name, section_rows in sections:
                if not section_rows and section_name:
                    footer_lines.append(section_name)
                else:
                    live_sections.append((section_name, section_rows))
            if footer_lines:
                pre_header_lines.extend(footer_lines)
            sections = live_sections

            # Step 7: Convert each section into a markdown table
            for section_name, section_rows in sections:
                if not section_rows:
                    continue

                header_line = "| " + " | ".join(headers) + " |"
                separator = "| " + " | ".join(["---"] * len(headers)) + " |"

                md_lines = [header_line, separator]
                for row in section_rows:
                    cells = []
                    for c in row:
                        val = str(c).strip() if c is not None else ""
                        val = val.replace("\n", " ").replace("|", "/")
                        # Collapse vertically-written text like
                        # "F I N A N C I A L" -> "FINANCIAL"
                        if (len(val) >= 5
                                and all(
                                    (ch.isalpha() and len(ch) == 1)
                                    or ch == " "
                                    for ch in val)
                                and " " in val):
                            parts = val.split()
                            if all(len(p) == 1 for p in parts):
                                val = "".join(parts)
                        cells.append(val)
                    md_lines.append("| " + " | ".join(cells) + " |")

                table_text = "\n".join(md_lines)

                # Prepend sheet/section name so it is searchable
                sn_stripped = sheet_name.strip()
                if section_name:
                    table_text = (
                        f"[{sn_stripped} > {section_name}]\n" + table_text
                    )
                else:
                    table_text = f"[{sn_stripped}]\n" + table_text

                # Prepend pre-header context (title, definitions, etc.)
                if pre_header_lines:
                    context_block = "\n".join(pre_header_lines)
                    table_text = context_block + "\n\n" + table_text

                heading = section_name if section_name else sheet_name
                parent = (
                    f"{filename} > {sheet_name}"
                    if section_name
                    else ""
                )

                content.append({
                    "text": table_text,
                    "source_file": filename,
                    "sheet_name": sheet_name,
                    "doc_type": "tabular",
                    "heading_level": 0,
                    "heading_text": heading,
                    "parent_path": parent,
                })

            if not sections:
                logger.info(f"  Excel: sheet '{sheet_name}' produced no sections after parsing")

        wb.close()

        if content:
            total_sections = len(content)
            total_rows = sum(c["text"].count("\n") - 1 for c in content)
            logger.info(
                f"  Excel: {filename} -> {total_sections} section(s), "
                f"~{total_rows} data rows"
            )
        return content

    # ------------------------------------------------------------------
    # Excel helpers — merged cells, header detection, section splitting
    # ------------------------------------------------------------------

    @staticmethod
    def _xlsx_read_sheet_with_merges(ws) -> List[List[Any]]:
        """
        Read an openpyxl worksheet into a list-of-lists grid.

        Merged cells are forward-filled: the top-left value of each
        merged range is copied into every cell the range spans.
        """
        max_row = ws.max_row
        max_col = ws.max_column
        if not max_row or not max_col or max_row < 1:
            return []

        grid = []
        for row in ws.iter_rows(
            min_row=1, max_row=max_row, max_col=max_col, values_only=True
        ):
            grid.append(list(row))

        for merged_range in ws.merged_cells.ranges:
            min_r = merged_range.min_row - 1
            max_r = merged_range.max_row - 1
            min_c = merged_range.min_col - 1
            max_c = merged_range.max_col - 1
            value = grid[min_r][min_c] if min_r < len(grid) and min_c < len(grid[min_r]) else None
            for r in range(min_r, min(max_r + 1, len(grid))):
                for c in range(min_c, min(max_c + 1, len(grid[r]))):
                    grid[r][c] = value

        return grid

    @staticmethod
    def _xlsx_detect_header_row(grid: List[List[Any]]) -> Optional[int]:
        """
        Scan the first 15 rows to find the most likely header row.

        A header row has the highest count of cells that look like column
        names: non-null strings, length 2-80, not pure numbers or dates.
        Requires at least 3 qualifying cells.
        """
        best_idx = None
        best_score = 0
        scan_limit = min(15, len(grid))

        for row_idx in range(scan_limit):
            row = grid[row_idx]
            score = 0
            qualifying_texts = []
            for cell in row:
                if cell is None:
                    continue
                text = str(cell).strip()
                if not text or len(text) > 80:
                    continue
                if len(text) < 2 and not text.isalpha():
                    continue
                try:
                    float(text.replace(",", ""))
                    continue
                except ValueError:
                    pass
                if re.match(r"^\d{4}-\d{2}-\d{2}", text):
                    continue
                score += 1
                qualifying_texts.append(text)

            # Merged title rows have many cells but few unique values
            unique_count = len(set(qualifying_texts))
            if unique_count < 3:
                score = min(score, unique_count)

            if score >= 2 and score > best_score:
                best_score = score
                best_idx = row_idx

        return best_idx

    @staticmethod
    def _xlsx_split_sections(
        headers: List[str],
        data_rows: List[List[Any]],
    ) -> List[Tuple[Optional[str], List[List[Any]]]]:
        """
        Detect section separator rows and split data into named sections.

        A separator row has values in only the first 1-2 columns while
        the remaining columns are empty (e.g. "Corrugation - Related"
        spanning column A with columns B-G all null). The separator text
        becomes the section name; subsequent data rows belong to it.
        """
        if not data_rows or not headers:
            return [(None, data_rows)]

        num_cols = len(headers)
        if num_cols < 3:
            return [(None, data_rows)]

        sections: List[Tuple[Optional[str], List[List[Any]]]] = []
        current_name: Optional[str] = None
        current_rows: List[List[Any]] = []

        for row in data_rows:
            non_empty = [
                i for i, c in enumerate(row)
                if c is not None and str(c).strip()
            ]

            is_separator = (
                len(non_empty) <= 2
                and len(non_empty) >= 1
                and all(i < 2 for i in non_empty)
                and num_cols >= 4
            )

            if is_separator:
                sep_text = str(row[non_empty[0]]).strip()
                # Bare numbers (row counters like "1", "2") are data, not separators
                try:
                    float(sep_text.replace(",", ""))
                    current_rows.append(row)
                    continue
                except ValueError:
                    pass
                # Flush previous section (keep empty named sections for footer capture)
                if current_rows or current_name is not None:
                    sections.append((current_name, current_rows))
                current_name = sep_text
                current_rows = []
            else:
                current_rows.append(row)

        if current_rows or current_name is not None:
            sections.append((current_name, current_rows))

        return sections if sections else [(None, data_rows)]

    def _load_xls_bytes(self, data: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Legacy .xls extraction.

        Converts to in-memory .xlsx via pandas+xlrd, then delegates to
        the smart _load_xlsx_bytes parser. Falls back to raw xlrd cell
        extraction if conversion fails.
        """
        try:
            return self._load_xlsx_bytes(data, filename)
        except Exception:
            pass

        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=data)
            content = []
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                if sheet.nrows < 2:
                    continue
                rows = []
                for row_idx in range(sheet.nrows):
                    cells = []
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell(row_idx, col_idx)
                        cells.append(str(cell.value).strip())
                    rows.append(" | ".join(cells))
                sheet_text = "\n".join(rows)
                if sheet_text.strip():
                    content.append({
                        "text": sheet_text, "source_file": filename,
                        "sheet_name": sheet.name, "doc_type": "tabular",
                        "heading_level": 0, "heading_text": "", "parent_path": "",
                    })
            return content
        except ImportError:
            raise ValueError(
                f"Cannot read .xls file '{filename}'. "
                f"Install xlrd: pip install xlrd"
            )

    # ===================================================================
    # Shared helpers
    # ===================================================================

    @staticmethod
    def _extract_docx_table(table) -> str:
        """
        Convert a python-docx Table to a pipe-delimited markdown string.

        Handles merged cells (rowspan / colspan) correctly.

        python-docx exposes merged cells by returning the SAME underlying
        <w:tc> XML element object for every logical cell that belongs to the
        merge group.  Iterating ``row.cells`` therefore gives duplicate
        objects — e.g. a cell that spans 3 rows appears 3 times in each of
        those rows.  We deduplicate by tracking the identity (``id()``) of
        the underlying ``_tc`` element so each physical cell is emitted only
        once per row, while empty-string placeholders preserve column count
        for downstream markdown alignment.
        """
        rows_out = []
        for row in table.rows:
            seen_tc_ids: set = set()
            cells_out: list = []
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    # Merged continuation — emit an empty placeholder so
                    # column count stays consistent across all rows.
                    cells_out.append("")
                else:
                    seen_tc_ids.add(tc_id)
                    cells_out.append(cell.text.strip())
            rows_out.append(" | ".join(cells_out))
        return "\n".join(rows_out)

    # ===================================================================
    # Chunking — Parent-Child Aware
    # ===================================================================

    def chunk_document(
        self,
        document_content: List[Dict[str, Any]],
        chunk_size: int = CHUNK_SIZE,
        overlap: int = CHUNK_OVERLAP,
    ) -> List[Dict[str, Any]]:
        """Chunk with parent-child awareness for hierarchical content."""
        chunks = []

        # --- Filter binary/garbage sections before anything else ---
        # .doc files (and occasionally PDFs) can produce sections with mostly
        # non-ASCII bytes when the binary parser misreads the file structure.
        # These sections produce useless chunks and inflate the index.
        clean_content = [s for s in document_content if _is_valid_section(s.get("text", ""))]
        filtered = len(document_content) - len(clean_content)
        if filtered:
            logger.info(f"  Chunking: filtered {filtered} garbage section(s) (non-ASCII content)")

        hierarchical_sections = [s for s in clean_content if s.get("heading_level", 0) > 0]
        flat_sections = [s for s in clean_content if s.get("heading_level", 0) == 0]

        if hierarchical_sections:
            chunks.extend(self._chunk_hierarchical(hierarchical_sections, chunk_size, overlap))

        for doc_part in flat_sections:
            text = doc_part["text"]
            # Skip garbage / binary / too-short chunks before indexing
            if not _is_valid_section(text):
                continue
            doc_type = doc_part.get("doc_type", "narrative")

            if doc_type == "tabular":
                # --- Split large tables by ROW boundaries, not word count ---
                # This preserves markdown table structure in each chunk.
                table_lines = text.split('\n')
                # Check if this is a markdown table (has | separators)
                is_markdown_table = any(line.strip().startswith('|') for line in table_lines)

                if is_markdown_table and len(text.split()) > chunk_size:
                    # Split into sub-tables, keeping header in each
                    header_lines = []
                    data_lines = []
                    for line in table_lines:
                        stripped = line.strip()
                        if not stripped:
                            continue
                        if stripped.startswith('|') and all(c in '|- \t' for c in stripped):
                            # This is a separator line (|---|---|)
                            if not header_lines:
                                # Haven't seen header yet, skip separator for now
                                continue
                            # Separator after header — we already have the header
                            continue
                        elif not header_lines:
                            header_lines.append(line)
                        else:
                            data_lines.append(line)

                    # Build the separator line
                    if header_lines:
                        num_cols = header_lines[0].count('|') - 1
                        separator = '|' + '|'.join(['---'] * num_cols) + '|'
                    else:
                        separator = ''

                    # Split data rows into chunks that fit within chunk_size
                    current_rows = []
                    current_words = len(header_lines[0].split()) if header_lines else 0
                    header_word_count = current_words + len(separator.split()) if separator else current_words

                    for row in data_lines:
                        row_words = len(row.split())
                        if current_words + row_words > chunk_size and current_rows:
                            # Flush current chunk
                            chunk_text = '\n'.join(header_lines + [separator] + current_rows)
                            chunk = doc_part.copy()
                            chunk["text"] = chunk_text
                            chunk["chunk_id"] = len(chunks)
                            chunk["chunk_role"] = "standalone"
                            chunk["doc_type"] = "tabular"
                            # Mark as part of a split table
                            chunks.append(chunk)
                            # Start new chunk with header repeated
                            current_rows = [row]
                            current_words = header_word_count + row_words
                        else:
                            current_rows.append(row)
                            current_words += row_words

                        # If a single row exceeds chunk_size * 3, split it
                        # with a word-based sliding window (wide guidance tables
                        # can have 4000+ words in one row).
                        if (len(current_rows) == 1
                                and len(current_rows[0].split()) > chunk_size * 3):
                            oversized = current_rows[0]
                            ow = oversized.split()
                            current_rows = []
                            for wi in range(0, len(ow), chunk_size - overlap):
                                sub = " ".join(ow[wi: wi + chunk_size])
                                if sub:
                                    ct = '\n'.join(
                                        header_lines + [separator] + [sub])
                                    ch = doc_part.copy()
                                    ch["text"] = ct
                                    ch["chunk_id"] = len(chunks)
                                    ch["chunk_role"] = "standalone"
                                    ch["doc_type"] = "tabular"
                                    chunks.append(ch)
                            current_words = header_word_count

                    # Flush remaining rows
                    if current_rows:
                        chunk_text = '\n'.join(header_lines + [separator] + current_rows)
                        chunk = doc_part.copy()
                        chunk["text"] = chunk_text
                        chunk["chunk_id"] = len(chunks)
                        chunk["chunk_role"] = "standalone"
                        chunk["doc_type"] = "tabular"
                        chunks.append(chunk)
                else:
                    # Small table — keep as single chunk
                    chunk = doc_part.copy()
                    chunk["chunk_id"] = len(chunks)
                    chunk["chunk_role"] = "standalone"
                    chunks.append(chunk)
            else:
                words = text.split()
                if len(words) <= chunk_size:
                    chunk = doc_part.copy()
                    chunk["chunk_id"] = len(chunks)
                    chunk["chunk_role"] = "standalone"
                    chunks.append(chunk)
                else:
                    for i in range(0, len(words), chunk_size - overlap):
                        chunk_text = " ".join(words[i : i + chunk_size])
                        if chunk_text:
                            chunk = doc_part.copy()
                            chunk["text"] = chunk_text
                            chunk["chunk_id"] = len(chunks)
                            chunk["chunk_role"] = "standalone"
                            chunks.append(chunk)

        return chunks

    def _chunk_hierarchical(
        self, sections: List[Dict[str, Any]], chunk_size: int, overlap: int,
    ) -> List[Dict[str, Any]]:
        """
        Build parent chunks from heading-based sections.

        Each top-level section becomes ONE parent chunk whose text is the
        heading + all child section text concatenated. Child sections are
        NOT stored as separate chunks — the parent already contains everything.

        This eliminates the duplicate "parent + children" problem where the
        same content was sent to the LLM twice (once in the parent, once in
        each child). The retrieval layer uses the parent for both embedding
        search and for LLM context.

        If the parent text exceeds chunk_size words, it is split into
        overlapping standalone chunks so no content is lost.
        """
        chunks = []
        if not sections:
            return chunks

        min_level = min(s["heading_level"] for s in sections)
        groups = self._group_by_top_heading(sections, min_level)

        for group in groups:
            parent_section = group["parent"]
            child_sections = group["children"]

            # Filter garbage child sections before merging
            valid_children = [c for c in child_sections if _is_valid_section(c["text"])]

            # Build full section text: parent heading text + all child text
            text_parts = [parent_section["text"]]
            for child_sec in valid_children:
                text_parts.append(child_sec["text"])
            full_text = "\n\n".join(text_parts)

            # Skip if the combined section is garbage
            if not _is_valid_section(full_text):
                continue

            words = full_text.split()

            if len(words) <= chunk_size:
                # Fits in one chunk — single standalone chunk with full section text
                parent_chunk = {
                    "text": full_text,
                    "source_file": parent_section["source_file"],
                    "doc_type": "narrative",
                    "heading_level": parent_section["heading_level"],
                    "heading_text": parent_section["heading_text"],
                    "parent_path": parent_section["parent_path"],
                    "chunk_id": len(chunks),
                    "chunk_role": "standalone",
                    "parent_id": None,
                    "child_ids": [],
                }
                if "page_num" in parent_section:
                    parent_chunk["page_num"] = parent_section["page_num"]
                chunks.append(parent_chunk)

            else:
                # Section is large — split into overlapping chunks.
                # ALL split parts keep the heading metadata so every part is
                # retrievable when someone asks about this section.
                parts = []
                for i in range(0, len(words), chunk_size - overlap):
                    chunk_text = " ".join(words[i: i + chunk_size])
                    if chunk_text:
                        parts.append(chunk_text)

                for part_idx, chunk_text in enumerate(parts):
                    chunk = {
                        "text": chunk_text,
                        "source_file": parent_section["source_file"],
                        "doc_type": "narrative",
                        # Keep heading on ALL parts — ensures every split is
                        # retrievable when the query mentions this section name
                        "heading_level": parent_section["heading_level"],
                        "heading_text": parent_section["heading_text"],
                        "parent_path": parent_section["parent_path"],
                        "chunk_id": len(chunks),
                        "chunk_role": "standalone",
                        "parent_id": None,
                        "child_ids": [],
                        # Mark as part N of M so LLM knows it's a continuation
                        "section_part": part_idx + 1,
                        "section_total_parts": len(parts),
                    }
                    if "page_num" in parent_section:
                        chunk["page_num"] = parent_section["page_num"]
                    chunks.append(chunk)

        return chunks

    def _group_by_top_heading(
        self, sections: List[Dict[str, Any]], top_level: int
    ) -> List[Dict[str, Any]]:
        """Group sections by their top-level heading."""
        groups = []
        current_parent = None
        current_children = []

        for sec in sections:
            if sec["heading_level"] == top_level:
                if current_parent is not None:
                    groups.append({"parent": current_parent, "children": current_children})
                current_parent = sec
                current_children = []
            else:
                current_children.append(sec)

        if current_parent is not None:
            groups.append({"parent": current_parent, "children": current_children})

        return groups


# ===========================================================================
# RAG System — Production Ready (Ollama + Parent-Child + Incremental + IVF)
# ===========================================================================
class RAGSystem:
    """
    Production-ready RAG pipeline — hybrid Python+LibreOffice, no OpenAI:
      - Ingest from Azure Blob Storage (streaming) or local files
      - PARALLEL blob streaming with configurable workers
      - INCREMENTAL ingestion with blob manifest tracking
      - PROGRESS CHECKPOINTING — auto-save every N documents
      - Hierarchical parent-child chunking for DOCX, DOC, PDF, and ODT
      - PDF font-size based heading detection + heuristic fallback
      - .doc extraction via LibreOffice conversion to DOCX + hierarchical parsing
      - HuggingFace embeddings — no external API needed
      - FAISS IVF index for fast search at 2K+ scale (auto-upgrades from flat)
      - BM25 sparse index + Hybrid RRF retrieval
      - Parent-child expansion for full context
      - Ollama LLM for generation with retry logic
      - chunks_output.txt auto-generated after ingestion
      - Proper logging with timestamps
    """

    def __init__(
        self,
        azure_conn_str: Optional[str] = None,
        document_container: Optional[str] = None,
        vector_container: Optional[str] = None,
        ollama_base_url: Optional[str] = None,
        model_name: Optional[str] = None,
        embedding_model: Optional[str] = None,
        faiss_local_path: Optional[str] = None,
        chunks_output_file: Optional[str] = None,
        debug_chunks_file: Optional[str] = None,
        ssl_verify: Optional[bool] = None,
        parallel_workers: Optional[int] = None,
        checkpoint_every: Optional[int] = None,
        embedding_batch_size: Optional[int] = None,
        incremental: bool = True,
    ):
        self.azure_conn_str       = azure_conn_str or AZURE_STORAGE_CONNECTION_STRING
        self.document_container   = document_container or DOCUMENT_CONTAINER
        self.vector_container     = vector_container or VECTOR_CONTAINER
        self.ollama_base_url      = ollama_base_url or OLLAMA_BASE_URL
        self.model_name           = model_name or MODEL_NAME
        self.embedding_model_name = embedding_model or EMBEDDING_MODEL
        self.faiss_local_path     = faiss_local_path or FAISS_LOCAL_PATH
        self.chunks_output_file   = chunks_output_file or CHUNKS_OUTPUT_FILE
        self.debug_chunks_file    = debug_chunks_file if debug_chunks_file is not None else DEBUG_CHUNKS_FILE
        self.ssl_verify           = ssl_verify if ssl_verify is not None else SSL_VERIFY
        self.parallel_workers     = parallel_workers or PARALLEL_WORKERS
        self.checkpoint_every     = checkpoint_every or CHECKPOINT_EVERY
        self.embedding_batch_size = embedding_batch_size or EMBEDDING_BATCH_SIZE
        self.incremental          = incremental
        self.ivf_nlist            = FAISS_IVF_NLIST
        self.ivf_threshold        = FAISS_IVF_THRESHOLD
        self.nprobe               = FAISS_NPROBE

        logger.info(f"Loading embedding model: {self.embedding_model_name} ...")
        self.embedder = SentenceTransformer(self.embedding_model_name)

        self._blob_manager = None
        self._vector_blob_manager = None
        self.chunk_writer = ChunkOutputWriter(self.chunks_output_file)

        self.document_processor = DocumentProcessor()
        self.vector_store = None
        self._index_type = "flat"  # Track whether we're using flat or IVF
        self.documents: List[Dict[str, Any]] = []
        self.bm25 = None
        self._file_bm25 = None
        self._file_names_ordered = []
        self._file_chunk_groups = {}

        # Ingestion manifest for incremental processing
        manifest_path = os.path.join(self.faiss_local_path, "ingested_manifest.json")
        self.manifest = IngestionManifest(manifest_path)

        # Stats tracking
        self._ingestion_stats = {
            "total_docs": 0,
            "total_chunks": 0,
            "failed_docs": 0,
            "start_time": None,
            "end_time": None,
        }
        # Tracks (file_name_or_path, error_reason) for every doc that failed ingestion
        self._failed_doc_details: List[Tuple[str, str]] = []
        self._last_retrieval_debug = None

    @property
    def blob_manager(self) -> Optional[AzureBlobManager]:
        if not self.azure_conn_str:
            return None
        if self._blob_manager is None:
            self._blob_manager = AzureBlobManager(self.azure_conn_str, self.document_container)
        return self._blob_manager

    @property
    def vector_blob_manager(self) -> Optional[AzureBlobManager]:
        if not self.azure_conn_str:
            return None
        if self._vector_blob_manager is None:
            self._vector_blob_manager = AzureBlobManager(self.azure_conn_str, self.vector_container)
        return self._vector_blob_manager

    # ------------------------------------------------------------------
    # Ingestion Stats
    # ------------------------------------------------------------------
    def _write_failed_docs_report(self):
        """Write a .txt report listing every document that failed ingestion."""
        report_path = os.path.join(self.faiss_local_path, "failed_documents.txt")
        try:
            os.makedirs(self.faiss_local_path, exist_ok=True)
            with open(report_path, "w", encoding="utf-8") as f:
                f.write("FAILED DOCUMENTS REPORT\n")
                f.write("=" * 60 + "\n")
                f.write(f"Generated: {datetime.now().isoformat()}\n")
                f.write(f"Total failed: {len(self._failed_doc_details)}\n")
                f.write("=" * 60 + "\n\n")

                if not self._failed_doc_details:
                    f.write("No failed documents. All documents ingested successfully.\n")
                else:
                    for i, (doc_name, reason) in enumerate(self._failed_doc_details, 1):
                        f.write(f"{i}. {doc_name}\n")
                        f.write(f"   Reason: {reason}\n\n")

            logger.info(f"  Failed docs report written: {report_path}")
        except Exception as e:
            logger.error(f"Could not write failed docs report: {e}")

    def _print_stats(self):
        """Print ingestion summary."""
        stats = self._ingestion_stats
        elapsed = 0
        if stats["start_time"] and stats["end_time"]:
            elapsed = (stats["end_time"] - stats["start_time"]).total_seconds()

        logger.info("=" * 60)
        logger.info("  INGESTION SUMMARY")
        logger.info("=" * 60)
        logger.info(f"  Documents processed:  {stats['total_docs']}")
        logger.info(f"  Documents failed:     {stats['failed_docs']}")
        logger.info(f"  Total chunks created: {stats['total_chunks']}")
        logger.info(f"  Total vectors in FAISS: {self.vector_store.ntotal if self.vector_store else 0}")
        logger.info(f"  FAISS index type:     {self._index_type}")
        logger.info(f"  Elapsed time:         {elapsed:.1f}s")
        logger.info(f"  Manifest entries:     {self.manifest.count}")
        logger.info("=" * 60)

        # Always write the failed-docs report (even if empty) so the user
        # has a single place to check after every ingestion run.
        self._write_failed_docs_report()

    # ------------------------------------------------------------------
    # Parallel Blob Streaming
    # ------------------------------------------------------------------
    def _stream_blobs_parallel(
        self, blob_names: List[str], max_workers: int
    ) -> Dict[str, bytes]:
        """Download multiple blobs in parallel using thread pool."""
        results = {}

        def _download_one(blob_name: str) -> Tuple[str, bytes]:
            data = self.blob_manager.stream_blob(blob_name)
            return blob_name, data

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(_download_one, name): name for name in blob_names}

            for future in as_completed(futures):
                blob_name = futures[future]
                try:
                    name, data = future.result()
                    results[name] = data
                except Exception as e:
                    logger.error(f"Failed to download blob '{blob_name}': {e}")
                    self._ingestion_stats["failed_docs"] += 1
                    self._failed_doc_details.append((blob_name, f"Download failed: {e}"))

        return results

    # ------------------------------------------------------------------
    # Blob Ingestion — with Incremental + Parallel + Checkpointing
    # ------------------------------------------------------------------
    def ingest_from_blob(self, prefix: Optional[str] = None):
        if not self.blob_manager:
            raise ValueError("AZURE_STORAGE_CONNECTION_STRING not configured.")

        self._ingestion_stats["start_time"] = datetime.now()

        # Step 1: Load existing index if incremental mode
        if self.incremental and self.vector_store is None:
            loaded = self.load_index()
            if loaded:
                logger.info(f"Loaded existing index: {self.vector_store.ntotal} vectors, "
                          f"{len(self.documents)} chunks, {self.manifest.count} tracked docs")

        # Step 2: List blobs and filter to supported types
        blob_names = self.blob_manager.list_blobs(prefix=prefix)
        supported_exts = set(DocumentProcessor.SUPPORTED_EXTENSIONS.keys())
        blob_names = [b for b in blob_names if any(b.lower().endswith(ext) for ext in supported_exts)]

        if not blob_names:
            logger.info("No supported documents found in Azure Blob container.")
            return

        # Step 3: Filter out already-ingested blobs (incremental)
        if self.incremental:
            new_blobs = self.manifest.get_unprocessed(blob_names)
            skipped = len(blob_names) - len(new_blobs)
            if skipped > 0:
                logger.info(f"Incremental mode: skipping {skipped} already-ingested documents")
            blob_names = new_blobs

        if not blob_names:
            logger.info("No new documents to ingest.")
            return

        logger.info(f"Found {len(blob_names)} new document(s) to ingest from Azure Blob.")

        # Step 4: Process documents in batches with parallel downloading
        all_chunks = []
        docs_processed = 0
        batch_size = self.parallel_workers * 2  # Download in batches

        for batch_start in range(0, len(blob_names), batch_size):
            batch_blobs = blob_names[batch_start:batch_start + batch_size]
            logger.info(f"Downloading batch {batch_start // batch_size + 1}/"
                       f"{(len(blob_names) + batch_size - 1) // batch_size} "
                       f"({len(batch_blobs)} blobs)...")

            # Parallel download
            blob_data = self._stream_blobs_parallel(batch_blobs, self.parallel_workers)

            # Process each downloaded blob
            for blob_name, data in blob_data.items():
                try:
                    doc_content = self.document_processor.load_from_bytes(
                        data, os.path.basename(blob_name)
                    )
                    chunks = self.document_processor.chunk_document(doc_content)

                    # Re-index chunks to global IDs
                    for chunk in chunks:
                        chunk["chunk_id"] = len(self.documents) + len(all_chunks)
                        all_chunks.append(chunk)

                    # Track in manifest
                    total_chars = sum(len(c.get("text", "")) for c in chunks)
                    self.manifest.mark_ingested(blob_name, len(chunks), total_chars)

                    docs_processed += 1
                    self._ingestion_stats["total_docs"] += 1
                    logger.info(f"  [{docs_processed}/{len(blob_names)}] {blob_name} "
                               f"-> {len(chunks)} chunks")

                except Exception as e:
                    logger.error(f"  Error processing {blob_name}: {e}")
                    self._ingestion_stats["failed_docs"] += 1
                    self._failed_doc_details.append((blob_name, str(e)))

            # Step 5: Checkpoint — save progress periodically
            if all_chunks and docs_processed % self.checkpoint_every == 0:
                logger.info(f"Checkpoint: saving progress ({docs_processed} docs, "
                           f"{len(all_chunks)} new chunks)...")
                self._add_chunks(all_chunks)
                self.chunk_writer.write_chunks(all_chunks)
                all_chunks = []  # Reset after saving
                self.save_index()
                self.manifest.save()

        # Step 6: Add remaining chunks
        if all_chunks:
            self._add_chunks(all_chunks)
            self.chunk_writer.write_chunks(all_chunks)

        # Step 7: Final save (may upgrade to IVF if threshold met)
        self.save_index()
        self.manifest.save()
        self.chunk_writer.save_to_file()

        self._ingestion_stats["end_time"] = datetime.now()
        self._ingestion_stats["total_chunks"] = len(self.documents)
        self._print_stats()

    def ingest_documents(self, file_paths: List[str]):
        """Ingest from local file paths with incremental support."""
        self._ingestion_stats["start_time"] = datetime.now()

        # Load existing index if incremental mode
        if self.incremental and self.vector_store is None:
            loaded = self.load_index()
            if loaded:
                logger.info(f"Loaded existing index: {self.vector_store.ntotal} vectors, "
                          f"{len(self.documents)} chunks")

        all_chunks = []
        docs_processed = 0

        for file_path in file_paths:
            if os.path.isdir(file_path):
                for root, _, files in os.walk(file_path):
                    for file in files:
                        full_path = os.path.join(root, file)
                        # Incremental: check manifest
                        if self.incremental and self.manifest.is_ingested(full_path):
                            continue
                        if any(file.lower().endswith(ext) for ext in DocumentProcessor.SUPPORTED_EXTENSIONS):
                            file_chunks = self._process_single_file(full_path)
                            if file_chunks:
                                for chunk in file_chunks:
                                    chunk["chunk_id"] = len(self.documents) + len(all_chunks)
                                    all_chunks.append(chunk)
                                self.manifest.mark_ingested(
                                    full_path, len(file_chunks),
                                    sum(len(c.get("text", "")) for c in file_chunks)
                                )
                                docs_processed += 1
                                self._ingestion_stats["total_docs"] += 1
                                logger.info(f"  [{docs_processed}] {full_path} -> {len(file_chunks)} chunks")

                                # Checkpoint
                                if docs_processed % self.checkpoint_every == 0:
                                    logger.info(f"Checkpoint: saving progress...")
                                    self._add_chunks(all_chunks)
                                    self.chunk_writer.write_chunks(all_chunks)
                                    all_chunks = []
                                    self.save_index()
                                    self.manifest.save()
            else:
                # Single file
                if self.incremental and self.manifest.is_ingested(file_path):
                    logger.info(f"Skipping already-ingested: {file_path}")
                    continue

                file_chunks = self._process_single_file(file_path)
                if file_chunks:
                    for chunk in file_chunks:
                        chunk["chunk_id"] = len(self.documents) + len(all_chunks)
                        all_chunks.append(chunk)
                    self.manifest.mark_ingested(
                        file_path, len(file_chunks),
                        sum(len(c.get("text", "")) for c in file_chunks)
                    )
                    docs_processed += 1
                    self._ingestion_stats["total_docs"] += 1
                    logger.info(f"  {file_path} -> {len(file_chunks)} chunks")

        if not all_chunks:
            logger.info("No new documents found to ingest.")
            return

        self._add_chunks(all_chunks)
        self.chunk_writer.write_chunks(all_chunks)
        self.save_index()
        self.manifest.save()
        self.chunk_writer.save_to_file()

        self._ingestion_stats["end_time"] = datetime.now()
        self._ingestion_stats["total_chunks"] = len(self.documents)
        self._print_stats()

    def _process_single_file(self, file_path: str) -> List[Dict[str, Any]]:
        try:
            doc_content = self.document_processor.load_document(file_path)
            return self.document_processor.chunk_document(doc_content)
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            self._ingestion_stats["failed_docs"] += 1
            self._failed_doc_details.append((file_path, str(e)))
            return []

    def _add_chunks(self, chunks: List[Dict[str, Any]]):
        """Add chunks to the system and rebuild indexes.

        Deduplicates by content hash so that re-ingesting the same file (or
        re-running ingestion on a corpus that was already partially indexed)
        never inserts duplicate chunks into the vector store.
        """
        # Build a hash set of every text already stored
        if not hasattr(self, "_chunk_hashes"):
            self._chunk_hashes: set = {
                _chunk_content_hash(d["text"]) for d in self.documents
            }

        new_unique: List[Dict[str, Any]] = []
        skipped = 0
        for chunk in chunks:
            h = _chunk_content_hash(chunk.get("text", ""))
            if h in self._chunk_hashes:
                skipped += 1
                continue
            self._chunk_hashes.add(h)
            new_unique.append(chunk)

        if skipped:
            logger.info(f"  Dedup: skipped {skipped} duplicate chunk(s), keeping {len(new_unique)} new.")

        if not new_unique:
            return

        self.documents.extend(new_unique)
        self._build_vector_store(new_unique)
        self._build_bm25_index()

    # ------------------------------------------------------------------
    # Embedding — with batch optimization
    # ------------------------------------------------------------------
    def _get_embedding(self, text: str) -> List[float]:
        return self.embedder.encode(text, normalize_embeddings=True).tolist()

    def _get_embeddings_batch(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings in optimized batches for large datasets."""
        all_embeddings = []
        total = len(texts)
        batch_size = self.embedding_batch_size

        for start in range(0, total, batch_size):
            end = min(start + batch_size, total)
            batch_texts = texts[start:end]
            embeddings = self.embedder.encode(
                batch_texts,
                normalize_embeddings=True,
                show_progress_bar=False,
                batch_size=min(32, len(batch_texts)),
            )
            all_embeddings.append(embeddings.astype("float32"))
            if total > 100 and end % 500 == 0:
                logger.info(f"  Embedding progress: {end}/{total} ({end*100//total}%)")

        return np.vstack(all_embeddings) if all_embeddings else np.array([], dtype="float32")

    # ------------------------------------------------------------------
    # FAISS Vector Store — with IVF auto-upgrade
    # ------------------------------------------------------------------
    def _build_vector_store(self, chunks: List[Dict[str, Any]]):
        logger.info(f"Generating embeddings for {len(chunks)} chunks "
                   f"(model={self.embedding_model_name}, batch={self.embedding_batch_size}) ...")
        texts = [c["text"] for c in chunks]
        embeddings = self._get_embeddings_batch(texts)
        if self.vector_store is None:
            self.vector_store = faiss.IndexFlatIP(embeddings.shape[1])
            self._index_type = "flat"
        self.vector_store.add(embeddings)
        logger.info(f"Vector store updated. Total vectors: {self.vector_store.ntotal} (type: {self._index_type})")

    def _maybe_upgrade_to_ivf(self):
        """
        Upgrade FAISS index from IndexFlatIP to IndexIVFFlat when the
        number of vectors exceeds the IVF threshold.

        IVF (Inverted File Index) partitions vectors into clusters (Voronoi cells).
        At search time, only the nprobe closest clusters are searched, giving
        sub-linear search time. This is essential for 2K+ documents.

        Parameters:
          - nlist: number of Voronoi cells (clusters). Rule of thumb: sqrt(N)
          - nprobe: number of cells to search at query time. Higher = more accurate but slower.
        """
        if self.vector_store is None:
            return

        ntotal = self.vector_store.ntotal
        if ntotal < self.ivf_threshold:
            logger.info(f"Vector count ({ntotal}) below IVF threshold ({self.ivf_threshold}), keeping flat index")
            return

        if self._index_type == "ivf":
            return  # Already using IVF

        logger.info(f"Upgrading FAISS index to IVF ({ntotal} vectors >= {self.ivf_threshold} threshold)...")

        try:
            d = self.vector_store.d  # dimension

            # Calculate optimal nlist: rule of thumb is sqrt(N), capped at 4096
            nlist = min(int(np.sqrt(ntotal)), 4096)
            nlist = max(nlist, 10)  # minimum 10 clusters

            # Get all existing vectors from the flat index
            all_vectors = faiss.rev_swig_ptr(
                self.vector_store.get_xb(), ntotal * d
            ).reshape(ntotal, d).copy()

            # Create IVF index
            quantizer = faiss.IndexFlatIP(d)
            ivf_index = faiss.IndexIVFFlat(quantizer, d, nlist, faiss.METRIC_INNER_PRODUCT)

            # Train IVF with existing vectors
            logger.info(f"Training IVF index with nlist={nlist}...")
            ivf_index.train(all_vectors)

            # Add all vectors to the IVF index
            ivf_index.add(all_vectors)
            ivf_index.nprobe = self.nprobe

            self.vector_store = ivf_index
            self._index_type = "ivf"
            logger.info(f"FAISS index upgraded to IVF: nlist={nlist}, nprobe={self.nprobe}, "
                       f"vectors={ivf_index.ntotal}")
        except Exception as e:
            logger.warning(f"Failed to upgrade to IVF index, keeping flat: {e}")
            # Keep the flat index as fallback

    def save_index(self):
        if self.vector_store is None:
            return

        # Check if we should upgrade to IVF before saving
        self._maybe_upgrade_to_ivf()

        os.makedirs(self.faiss_local_path, exist_ok=True)
        faiss_path = os.path.join(self.faiss_local_path, "index.faiss")
        meta_path = os.path.join(self.faiss_local_path, "metadata.pkl")
        faiss.write_index(self.vector_store, faiss_path)
        with open(meta_path, "wb") as f:
            pickle.dump({
                "documents": self.documents,
                "bm25": self.bm25,
                "index_type": self._index_type,
            }, f)
        logger.info(f"Index saved to {self.faiss_local_path}/ "
                   f"({self.vector_store.ntotal} vectors, {len(self.documents)} chunks, type={self._index_type})")
        if self.vector_blob_manager:
            try:
                self._upload_index_to_blob(faiss_path, meta_path)
            except Exception as e:
                logger.warning(f"Failed to upload index to Azure Blob: {e}")

    def load_index(self) -> bool:
        os.makedirs(self.faiss_local_path, exist_ok=True)
        faiss_path = os.path.join(self.faiss_local_path, "index.faiss")
        meta_path = os.path.join(self.faiss_local_path, "metadata.pkl")
        if not os.path.exists(faiss_path) and self.vector_blob_manager:
            try:
                self._download_index_from_blob(faiss_path, meta_path)
            except Exception as e:
                logger.warning(f"Failed to download index from Azure Blob: {e}")
        if not os.path.exists(faiss_path) or not os.path.exists(meta_path):
            return False
        try:
            self.vector_store = faiss.read_index(faiss_path)
            with open(meta_path, "rb") as f:
                meta = pickle.load(f)
            self.documents = meta["documents"]
            self.bm25 = meta["bm25"]
            self._index_type = meta.get("index_type", "flat")

            # Set nprobe for IVF indexes
            if self._index_type == "ivf" and hasattr(self.vector_store, 'nprobe'):
                self.vector_store.nprobe = self.nprobe

            self._build_file_level_index()

            logger.info(f"Index loaded ({self.vector_store.ntotal} vectors, "
                       f"{len(self.documents)} chunks, type={self._index_type})")
            return True
        except Exception as e:
            logger.error(f"Error loading index: {e}")
            return False

    def _upload_index_to_blob(self, faiss_path: str, meta_path: str):
        with open(faiss_path, "rb") as f:
            self.vector_blob_manager.upload_bytes("index.faiss", f.read())
        with open(meta_path, "rb") as f:
            self.vector_blob_manager.upload_bytes("metadata.pkl", f.read())
        logger.info("Index uploaded to Azure Blob.")

    def _download_index_from_blob(self, faiss_path: str, meta_path: str):
        data = self.vector_blob_manager.download_bytes("index.faiss")
        with open(faiss_path, "wb") as f:
            f.write(data)
        data = self.vector_blob_manager.download_bytes("metadata.pkl")
        with open(meta_path, "wb") as f:
            f.write(data)
        logger.info("Index downloaded from Azure Blob.")

    # ------------------------------------------------------------------
    # BM25 Index
    # ------------------------------------------------------------------
    def _build_bm25_index(self):
        logger.info(f"Building BM25 index for {len(self.documents)} chunks...")
        tokenized_corpus = [chunk["text"].lower().split() for chunk in self.documents]
        self.bm25 = BM25Okapi(tokenized_corpus)
        logger.info("BM25 index built.")
        self._build_file_level_index()

    def _build_file_level_index(self):
        """Build a file-level BM25 index for cross-corpus source selection."""
        self._file_chunk_groups = {}
        for i, c in enumerate(self.documents):
            sf = c.get("source_file", "")
            if sf not in self._file_chunk_groups:
                self._file_chunk_groups[sf] = []
            self._file_chunk_groups[sf].append(i)

        self._file_names_ordered = sorted(self._file_chunk_groups.keys())
        file_docs = []
        for sf in self._file_names_ordered:
            chunk_ids = self._file_chunk_groups[sf]
            combined = sf + " " + " ".join(
                self.documents[i]["text"] for i in chunk_ids
            )
            file_docs.append(combined)

        file_tokenized = [d.lower().split() for d in file_docs]
        self._file_bm25 = BM25Okapi(file_tokenized)
        logger.info(
            f"File-level BM25 index built: {len(self._file_names_ordered)} files"
        )

    # ------------------------------------------------------------------
    # Retrieval — Parent-Child Aware Hybrid
    # ------------------------------------------------------------------
    # ------------------------------------------------------------------
    # Document name matching helpers — v2 (code-prefix/suffix stripping)
    # ------------------------------------------------------------------
    @staticmethod
    def _strip_code_prefix(stem: str) -> str:
        """Remove corporate document code prefixes from a filename stem.

        Examples:
          'mg-css-gl-it-001_use of ai notetaker in meeting' -> 'use of ai notetaker in meeting'
          'fs_prp_08 cleaning' -> 'cleaning'
          'hr-sop-07 attendance top-up policy' -> 'attendance top-up policy'
          'wi-05 sitich glue' -> 'sitich glue'
          'laptop policy' -> 'laptop policy' (no prefix to strip)
          'sop-dispatch' -> 'dispatch'
        """
        lower = stem.lower()
        m = _CODE_PREFIX_RE.match(lower)
        if m and m.end() < len(lower):
            stripped = lower[m.end():]
            if len(stripped.strip()) > 1:
                return stripped.strip()

        # Fallback patterns for prefixes the regex may miss
        if lower.startswith("sop"):
            rest = re.sub(r"^sop[-_\s]*", "", lower)
            if rest and len(rest.strip()) > 1:
                return rest.strip()

        if lower.startswith("wi-"):
            rest = re.sub(r"^wi[-_\s]*\d*\s*", "", lower)
            if rest and len(rest.strip()) > 1:
                return rest.strip()

        return lower

    @staticmethod
    def _strip_code_suffix(stem: str) -> str:
        """Remove trailing codes and dates from a filename stem.

        Examples:
          'risk assessment of suppliers_ fs.sp 4.00_ra-01 28.02.23' -> 'risk assessment of suppliers'
          'sop for intercompany tramsactions v1' -> 'sop for intercompany tramsactions'
        """
        result = _CODE_SUFFIX_RE.sub("", stem)
        result = result.strip().rstrip('_').rstrip('-').rstrip()
        return result if len(result) > 1 else stem

    def _get_meaningful_name(self, raw_stem: str) -> str:
        """Strip both code prefix AND suffix to get the meaningful document name."""
        return self._strip_code_suffix(self._strip_code_prefix(raw_stem))

    @staticmethod
    def _get_meaningful_words(text: str) -> set:
        """Extract meaningful (non-filler) words from text."""
        words = set(text.replace("_", " ").replace("-", " ").replace(",", " ").split())
        return words - _FILLER_WORDS

    @staticmethod
    def _clean_query(query: str) -> str:
        """Strip filler words and punctuation from a query for matching."""
        q = query.lower().replace("?", "").replace(".", "").replace("!", "")
        q = q.replace(",", "").replace(";", "").replace(":", "")
        words = q.split()
        meaningful = [w for w in words if w not in _FILLER_WORDS]
        return " ".join(meaningful)

    @staticmethod
    def _words_overlap(w1: str, w2: str) -> bool:
        """Check if two words overlap (handles plurals, tense variations).

        'supplier' and 'suppliers' -> True
        'assess' and 'assessment' -> True
        'calibrates' and 'calibration' -> True
        'laptop' and 'collaboration' -> False
        """
        if w1 == w2:
            return True
        shorter, longer = (w1, w2) if len(w1) <= len(w2) else (w2, w1)
        if longer.startswith(shorter) and len(shorter) >= 4:
            return True
        if len(shorter) >= 6:
            prefix_len = 0
            for c1, c2 in zip(w1, w2):
                if c1 == c2:
                    prefix_len += 1
                else:
                    break
            if prefix_len >= 6:
                return True
        return False

    @classmethod
    def _has_specific_overlap(cls, query_meaningful: set, stem_meaningful: set) -> bool:
        """Check if there's at least one NON-GENERIC meaningful keyword overlap.

        This prevents false positives like "collaboration tools policy"
        matching "laptop policy" — the only overlap ("policy") is generic.
        """
        for qw in query_meaningful:
            for sw in stem_meaningful:
                if cls._words_overlap(qw, sw):
                    if sw not in _GENERIC_DOC_WORDS and qw not in _GENERIC_DOC_WORDS:
                        return True
        # Also pass if 2+ keywords overlap in total (even if some are generic)
        overlap_count = 0
        for qw in query_meaningful:
            for sw in stem_meaningful:
                if cls._words_overlap(qw, sw):
                    overlap_count += 1
                    break
        return overlap_count >= 2

    def _match_doc_names(self, query: str) -> List[str]:
        """
        Check if the query mentions or closely matches any ingested document name.

        v2 matching strategy (5 methods tried in order):
          1. SUBSTRING (raw)   — raw filename stem found in query (or vice versa)
          2. SUBSTRING (clean) — clean name (code prefix/suffix stripped) in query
          3. KEYWORD           — ≥55% of meaningful stem words appear in the query
          4. FUZZY             — SequenceMatcher ratio ≥ 0.70 on clean name vs
                                 clean query (≥ 0.80 for short names ≤2 words)
          5. GENERIC_GUARD     — if 3+ DIFFERENT docs match, filter stays OFF

        Key improvements over v1:
          - Strips code prefixes (MG-CSS-GL-IT-003_, FS_PRP_08, WI-05)
          - Strips code suffixes (_FS.SP 4.00_RA-01 28.02.23, v1)
          - Removes filler words ("what is the") before matching
          - Handles plural/stem word overlap ("supplier" ≈ "suppliers")
          - Prevents false positives from generic words like "policy"
          - Allows matching all versions of the same document

        Examples:
          Query: "what is the password policy?"
            -> matches "MG-CSS-GL-IT-003_PASSWORD POLICY.pdf" (SUBSTRING clean)
          Query: "supplier risk assessment"
            -> matches all 3 RISK ASSESSMENT OF SUPPLIERS docs (KEYWORD)
          Query: "cleaning procedure"
            -> matches "FS_PRP_08 Cleaning.DOC" (SUBSTRING clean)
          Query: "policy"
            -> no match (too generic, hits GENERIC_GUARD)
        """
        if DOC_MATCH_THRESHOLD >= 1.0:
            return []

        # Collect unique source file names from the index
        source_files = list(set(d.get("source_file", "") for d in self.documents if d.get("source_file")))
        if not source_files:
            return []

        query_lower = query.lower()
        query_clean = self._clean_query(query)
        query_meaningful = self._get_meaningful_words(query_clean)

        matched_sources = []

        for src in source_files:
            # Extract the stem: "Laptop Policy.docx" -> "laptop policy"
            raw_stem = os.path.splitext(os.path.basename(src))[0].lower()

            # Strip code prefix and suffix to get the meaningful name
            clean_stem = self._get_meaningful_name(raw_stem)
            stem_meaningful = self._get_meaningful_words(clean_stem)

            # ---- Method 1: Substring match on raw stem (fastest) ----
            if raw_stem in query_lower or query_lower in raw_stem:
                matched_sources.append(src)
                continue

            # ---- Method 1b: Substring match on clean stem vs clean query ----
            if clean_stem in query_clean or query_clean in clean_stem:
                matched_sources.append(src)
                continue

            # ---- Method 2: Meaningful keyword overlap ----
            # Uses _words_overlap for fuzzy word matching (plurals, stems)
            if stem_meaningful and query_meaningful:
                overlap_words = set()
                for qw in query_meaningful:
                    for sw in stem_meaningful:
                        if self._words_overlap(qw, sw):
                            overlap_words.add(sw)
                            break

                if len(stem_meaningful) > 0 and len(overlap_words) >= 2:
                    overlap_ratio = len(overlap_words) / len(stem_meaningful)
                    if overlap_ratio >= MEANINGFUL_OVERLAP_THRESHOLD:
                        matched_sources.append(src)
                        continue

            # ---- Method 3: Fuzzy match on clean stem vs clean query ----
            best_ratio = 0.0

            ratio = SequenceMatcher(None, clean_stem, query_clean).ratio()
            best_ratio = max(best_ratio, ratio)

            query_words_list = query_clean.split()
            stem_words_list = clean_stem.split()

            if len(query_words_list) > 1:
                window_sizes = range(
                    max(2, min(len(stem_words_list), 2)),
                    min(len(query_words_list) + 1, len(stem_words_list) + 3)
                )
                for window_size in window_sizes:
                    for i in range(len(query_words_list) - window_size + 1):
                        window = " ".join(query_words_list[i:i + window_size])
                        ratio = SequenceMatcher(None, clean_stem, window).ratio()
                        best_ratio = max(best_ratio, ratio)

            # Adaptive threshold: short stems need higher fuzzy score
            threshold = FUZZY_THRESHOLD_SHORT if len(stem_meaningful) <= 2 else FUZZY_THRESHOLD

            if best_ratio >= threshold:
                # Secondary check: must have specific keyword overlap
                if self._has_specific_overlap(query_meaningful, stem_meaningful):
                    matched_sources.append(src)

        # ---- Generic query guard ----
        # If too many DIFFERENT documents match, the query is too generic
        if len(matched_sources) >= MAX_GENERIC_MATCHES:
            clean_names = set()
            for src in matched_sources:
                stem = os.path.splitext(os.path.basename(src))[0].lower()
                normalized = self._get_meaningful_name(stem)
                # Remove date-like patterns for version grouping
                normalized = re.sub(r'\s*\d{2}\.\d{2}(\.\d{2,4})?', '', normalized).strip()
                clean_names.add(normalized)

            if len(clean_names) > 1:
                # Multiple DIFFERENT documents matched -> too generic, keep filter OFF
                logger.info(
                    f"  Doc filter: GENERIC_GUARD triggered — "
                    f"{len(matched_sources)} docs matched across {len(clean_names)} different names -> filter OFF"
                )
                return []

        return matched_sources

    _STOP_WORDS = frozenset({
        'the','of','for','and','a','in','on','is','to','what','how','does',
        'are','by','from','with','do','this','that','an','at','it','its',
        'our','we','be','was','were','been','which','who','whom','where',
        'when','why','has','have','had',
    })
    _EXT_WORDS = frozenset({'xlsx','xls','xlsm','copy','doc','pdf','docx'})

    @staticmethod
    def _stem_word(w):
        if len(w) > 3 and w.endswith('s'):
            return w[:-1]
        return w

    @classmethod
    def _word_set(cls, text):
        return {cls._stem_word(w) for w in re.findall(r'\w+', text.lower())} - cls._STOP_WORDS

    @staticmethod
    def _get_source_group(fname):
        name = fname
        name = re.sub(r'^Copy of\s+', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s*-\s*Copy\b', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s*\(\d+\)', '', name)
        name = re.sub(r'\.(xlsx?|xlsm|pdf|docx?)$', '', name, flags=re.IGNORECASE)
        words = re.findall(r'[a-zA-Z]{2,}', name.lower())
        skip = {'copy','latest','update','new','old','final','revised',
                'version','philip','pc','of','for','the','and'}
        sig = [w for w in words if w not in skip]
        return ' '.join(sig[:4]) if sig else name.strip().lower()

    _EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xls", ".csv"}

    _DOMAIN_FILE_HINTS = [
        # (trigger_terms, file_substrings)
        # trigger_terms: if any term appears as substring in the lowercased query, this hint fires
        # file_substrings: matched against lowercased basename of source files
        (["calibrat", "instrument calibr"], ["calibration"]),
        (["kpi"], ["kpi"]),
        (["customer record", "customer code", "customer field", "customer template"], ["customer_template"]),
        (["pre-production", "pre production", "enters production",
          "job readiness", "before production"], ["pre-production", "pre production", "job readiness"]),
        (["supplier list", "suppliers list", "all suppliers",
          "ink supplier", "suppliers we purchase"], ["suppliers"]),
        (["prp ", "prerequisite prog"], ["haccp manual"]),
        (["utilities control", "utilities measure", "prp for utilities"], ["haccp manual"]),
        (["client information", "fsms client"], ["f101", "fsms client"]),
        (["food chain categor"], ["f101", "fsms client"]),
        (["document master", "doc master"], ["document master"]),
        (["retention period", "retention column"], ["document master"]),
        (["product description", "shelf life", "raw material spec"], ["rm_product", "product description"]),
    ]

    def _domain_keyword_fallback(self, query: str, source_type: str = "all") -> List[str]:
        """Fallback when _match_doc_names returns nothing: use domain-specific
        keyword hints to identify likely source file(s)."""
        query_lower = query.lower()
        source_files = list(set(
            d.get("source_file", "")
            for d in self.documents if d.get("source_file")
        ))

        matched = set()
        for triggers, file_patterns in self._DOMAIN_FILE_HINTS:
            triggered = any(t in query_lower for t in triggers)
            if not triggered:
                continue

            for sf in source_files:
                if source_type == "excel":
                    ext = os.path.splitext(sf)[1].lower()
                    if ext not in self._EXCEL_EXTENSIONS:
                        continue
                basename_lower = os.path.basename(sf).lower()
                if any(pat in basename_lower for pat in file_patterns):
                    matched.add(sf)

        return list(matched)

    def retrieve(self, query: str, k: int = TOP_K, source_type: str = "all") -> List[Dict[str, Any]]:
        """
        Retrieve top-k relevant chunks for the query.

        Dispatches to source-type-specific retrieval:
        - "all" (default): original retrieval logic for PDF/DOC content
        - "excel": enhanced retrieval with domain hints, cross-corpus
          scoring, rare-term injection, and chunk cap
        """
        if not self.vector_store or not self.bm25:
            return []
        if source_type == "excel":
            return self._retrieve_excel(query, k)
        return self._retrieve_default(query, k)

    # ------------------------------------------------------------------
    # Default retrieval — original logic for PDF/DOC content
    # ------------------------------------------------------------------
    def _retrieve_default(self, query: str, k: int = TOP_K) -> List[Dict[str, Any]]:
        """Original retrieval path: doc filter -> hybrid search -> dedup ->
        sibling/table expansion.  No domain hints, no cross-corpus scoring,
        no chunk cap."""

        # ---- Step 1: Document name filtering ----
        matched_docs = self._match_doc_names(query)
        doc_filter_active = len(matched_docs) > 0

        if doc_filter_active:
            logger.info(f"  Doc filter: query matches {matched_docs} -> only sending chunks from these doc(s)")
        else:
            logger.info(f"  Doc filter: no specific document matched -> searching all documents")

        # ---- Step 2: Hybrid search ----
        fetch_k = k * 5 if doc_filter_active else k * 3
        raw_results = self._retrieve_hybrid(query, fetch_k)

        seen_content_hashes: set = set()
        result: List[Dict[str, Any]] = []

        for chunk in raw_results:
            if doc_filter_active:
                if chunk.get("source_file", "") not in matched_docs:
                    continue

            h = _chunk_content_hash(chunk.get("text", ""))
            if h in seen_content_hashes:
                continue
            seen_content_hashes.add(h)
            result.append(chunk)
            if len(result) >= k:
                break

        # If doc filter matched but no chunks survived, fall back to unfiltered
        if not result and doc_filter_active:
            logger.info(
                f"  Doc filter fallback: no chunks from matched files — "
                f"falling back to unfiltered retrieval"
            )
            doc_filter_active = False
            seen_content_hashes.clear()
            for chunk in raw_results:
                h = _chunk_content_hash(chunk.get("text", ""))
                if h in seen_content_hashes:
                    continue
                seen_content_hashes.add(h)
                result.append(chunk)
                if len(result) >= k:
                    break

        # ---- Step 3: Sibling part expansion ----
        sibling_chunks = []
        for chunk in result:
            if chunk.get("section_total_parts", 1) <= 1:
                continue
            heading = chunk.get("heading_text", "")
            source  = chunk.get("source_file", "")
            total   = chunk.get("section_total_parts", 1)
            for doc in self.documents:
                if (doc.get("heading_text") == heading
                        and doc.get("source_file") == source
                        and doc.get("section_total_parts") == total):
                    h = _chunk_content_hash(doc.get("text", ""))
                    if h not in seen_content_hashes:
                        if doc_filter_active and doc.get("source_file", "") not in matched_docs:
                            continue
                        seen_content_hashes.add(h)
                        sibling_chunks.append(doc)

        sibling_chunks.sort(key=lambda c: c.get("section_part", 0))
        result.extend(sibling_chunks)

        # ---- Step 3b: Table expansion ----
        if doc_filter_active:
            tabular_sources_in_result = set()
            for chunk in result:
                if chunk.get("doc_type") == "tabular":
                    tabular_sources_in_result.add(chunk.get("source_file", ""))

            if tabular_sources_in_result:
                table_expansion_chunks = []
                for doc in self.documents:
                    if doc.get("doc_type") != "tabular":
                        continue
                    if doc.get("source_file", "") not in tabular_sources_in_result:
                        continue
                    h = _chunk_content_hash(doc.get("text", ""))
                    if h not in seen_content_hashes:
                        seen_content_hashes.add(h)
                        table_expansion_chunks.append(doc)
                result.extend(table_expansion_chunks)
                if table_expansion_chunks:
                    logger.info(f"  Table expansion: added {len(table_expansion_chunks)} additional table chunks")

        # ---- Step 4: Store retrieval debug info ----
        self._last_retrieval_debug = {
            "query": query,
            "doc_filter_active": doc_filter_active,
            "matched_docs": matched_docs,
            "raw_count": len(raw_results),
            "filtered_count": len(result),
            "sibling_count": len(sibling_chunks),
            "chunks": [
                {
                    "chunk_id": c.get("chunk_id"),
                    "source_file": c.get("source_file", ""),
                    "section": c.get("heading_text", ""),
                    "text": c.get("text", ""),
                    "page": c.get("page_num", ""),
                }
                for c in result
            ],
            "raw_results_preview": [
                {
                    "chunk_id": c.get("chunk_id"),
                    "source": c.get("source_file", "?"),
                    "heading": c.get("heading_text", ""),
                    "text_len": len(c.get("text", "")),
                    "page": c.get("page_num", ""),
                }
                for c in raw_results[:30]
            ],
        }

        logger.info(
            f"  Retrieval: {len(result)} chunks for LLM "
            f"({len(result) - len(sibling_chunks)} from search + {len(sibling_chunks)} sibling parts)"
            f"{' [DOC-FILTERED: ' + str(matched_docs) + ']' if doc_filter_active else ''}"
        )
        return result

    # ------------------------------------------------------------------
    # Excel retrieval — enhanced logic with all optimisations
    # ------------------------------------------------------------------
    def _retrieve_excel(self, query: str, k: int = TOP_K) -> List[Dict[str, Any]]:
        """Enhanced retrieval for Excel content: domain keyword fallback,
        family expansion, adaptive TOP_K, cross-corpus scoring, rare-term
        injection, BM25 direct scan, diversity cap, and chunk cap."""

        # ---- Step 1: Document name filtering (Excel only) ----
        all_matched = self._match_doc_names(query)
        matched_docs = [
            m for m in all_matched
            if os.path.splitext(m)[1].lower() in self._EXCEL_EXTENSIONS
        ]
        doc_filter_active = len(matched_docs) > 0
        logger.info(
            f"  Source type filter: EXCEL only"
            f"{' (doc filter: ' + str(matched_docs) + ')' if doc_filter_active else ''}"
        )

        # ---- Step 1a: Domain keyword fallback ----
        if not doc_filter_active:
            hint_matches = self._domain_keyword_fallback(query, "excel")
            if hint_matches:
                matched_docs = hint_matches
                doc_filter_active = True
                logger.info(
                    f"  Domain keyword fallback: {hint_matches}"
                )

        # ---- Step 1b: Doc filter family expansion ----
        if doc_filter_active:
            all_sources = list(set(
                d.get("source_file", "")
                for d in self.documents if d.get("source_file")
            ))
            expanded = set(matched_docs)
            for md in matched_docs:
                md_stem = os.path.splitext(os.path.basename(md))[0].lower()
                md_words = self._get_meaningful_words(
                    self._get_meaningful_name(md_stem)
                )
                if len(md_words) < 2:
                    continue
                for sf in all_sources:
                    if sf in expanded:
                        continue
                    ext = os.path.splitext(sf)[1].lower()
                    if ext not in self._EXCEL_EXTENSIONS:
                        continue
                    sf_stem = os.path.splitext(os.path.basename(sf))[0].lower()
                    sf_words = self._get_meaningful_words(
                        self._get_meaningful_name(sf_stem)
                    )
                    overlap = md_words & sf_words
                    if len(overlap) >= 2 and len(overlap) / len(md_words) >= 0.5:
                        expanded.add(sf)
            if len(expanded) > len(matched_docs):
                logger.info(
                    f"  Doc filter family expansion: {matched_docs} "
                    f"-> {sorted(expanded)}"
                )
                matched_docs = list(expanded)

        if doc_filter_active:
            logger.info(f"  Doc filter: query matches {matched_docs} -> only sending chunks from these doc(s)")
        else:
            logger.info(f"  Doc filter: no specific document matched -> searching all Excel documents")

        # ---- Step 1c: Adaptive TOP_K ----
        if doc_filter_active:
            effective_size = sum(
                1 for d in self.documents
                if d.get("source_file", "") in matched_docs
            )
        else:
            effective_size = len(self.documents)

        if effective_size <= 5:
            k = max(k, 5)
        elif effective_size <= 15:
            k = max(k, 8)
        elif effective_size <= 25:
            k = max(k, 10)
        elif effective_size <= 50:
            k = max(k, 15)
        elif effective_size <= 200:
            k = max(k, 15)
        elif effective_size <= 600:
            k = max(k, 20)
        else:
            k = max(k, 30)

        # ---- Step 1d: File-level source scoring (cross-corpus) ----
        file_boost_map = {}
        filename_matched = {}
        _excel_soft_boost = False
        _use_cross_corpus = (
            not doc_filter_active
            and self._file_bm25
            and len(self._file_names_ordered) >= 5
            and len(self.documents) >= 500
        )
        if _use_cross_corpus and self.bm25:
            q_tokens = query.lower().split()
            chunk_bm25 = self.bm25.get_scores(q_tokens)

            excel_file_top: Dict[str, List[float]] = {}
            for i, doc_chunk in enumerate(self.documents):
                sf = doc_chunk.get("source_file", "")
                ext = os.path.splitext(sf)[1].lower()
                if ext not in self._EXCEL_EXTENSIONS:
                    continue
                if sf not in excel_file_top:
                    excel_file_top[sf] = []
                excel_file_top[sf].append(chunk_bm25[i])

            file_agg = []
            for sf, scores in excel_file_top.items():
                top3 = sorted(scores, reverse=True)[:3]
                file_agg.append((sf, sum(top3)))
            file_agg.sort(key=lambda x: x[1], reverse=True)

            if file_agg and file_agg[0][1] > 0:
                max_agg = file_agg[0][1]
                for sf, sc in file_agg:
                    file_boost_map[sf] = 1.0 + 3.0 * (sc / max_agg)
                _excel_soft_boost = True
                logger.info(
                    f"  Excel soft boost (chunk-agg): top files "
                    f"{[(sf, round(sc, 1)) for sf, sc in file_agg[:5]]}"
                )

        # ---- Step 2: Hybrid search ----
        fetch_k = k * 10
        raw_results = self._retrieve_hybrid(query, fetch_k)

        raw_results = [
            c for c in raw_results
            if os.path.splitext(c.get("source_file", ""))[1].lower()
            in self._EXCEL_EXTENSIONS
        ]
        logger.info(f"  Excel filter: {len(raw_results)} chunks after filtering")

        # ---- Step 2b: Rare-term chunk injection ----
        if self.bm25 and not doc_filter_active:
            _stop = self._STOP_WORDS | _FILLER_WORDS
            q_words = {
                w for w in re.findall(r'\w+', query.lower())
                if len(w) >= 3 and w not in _stop
            }
            existing_ids = set(c.get("chunk_id") for c in raw_results)
            rare_injected = 0
            rare_threshold = max(int(len(self.documents) * 0.005), 5)

            for word in q_words:
                scores = self.bm25.get_scores([word])
                matching_idx = np.where(scores > 0)[0]

                if len(matching_idx) == 0 or len(matching_idx) > rare_threshold:
                    continue

                for ci in matching_idx:
                    if ci in existing_ids:
                        continue
                    doc = self.documents[ci]
                    ext = os.path.splitext(doc.get("source_file", ""))[1].lower()
                    if ext not in self._EXCEL_EXTENSIONS:
                        continue
                    raw_results.insert(0, doc)
                    existing_ids.add(ci)
                    rare_injected += 1

            if rare_injected:
                matched_terms = q_words & set(re.findall(r'\w+', query.lower()))

                logger.info(
                    f"  Rare-term injection: added {rare_injected} chunks "
                    f"for rare query terms {matched_terms}"
                )

                if _excel_soft_boost and file_boost_map:
                    max_boost = max(file_boost_map.values())
                    for chunk in raw_results[:rare_injected]:
                        sf = chunk.get("source_file", "")
                        if sf and file_boost_map.get(sf, 1.0) < max_boost:
                            file_boost_map[sf] = max_boost
                            logger.info(
                                f"  Rare-term boost: {sf} -> {max_boost:.1f}"
                            )

        seen_content_hashes: set = set()
        result: List[Dict[str, Any]] = []

        if doc_filter_active:
            for chunk in raw_results:
                if chunk.get("source_file", "") not in matched_docs:
                    continue
                h = _chunk_content_hash(chunk.get("text", ""))
                if h in seen_content_hashes:
                    continue
                seen_content_hashes.add(h)
                result.append(chunk)
                if len(result) >= k:
                    break

            if not result and self.bm25:
                logger.info(
                    f"  Doc filter matched {matched_docs} but no chunks in "
                    f"top-{fetch_k} hybrid — scanning matched file directly"
                )
                q_tokens = query.lower().split()
                bm25_scores = self.bm25.get_scores(q_tokens)
                candidates = []
                for i, doc_chunk in enumerate(self.documents):
                    if doc_chunk.get("source_file", "") in matched_docs:
                        candidates.append((bm25_scores[i], i, doc_chunk))
                candidates.sort(key=lambda x: x[0], reverse=True)
                for _sc, _i, chunk in candidates:
                    h = _chunk_content_hash(chunk.get("text", ""))
                    if h in seen_content_hashes:
                        continue
                    seen_content_hashes.add(h)
                    result.append(chunk)
                    if len(result) >= k:
                        break

            if not result:
                logger.info(
                    f"  Doc filter fallback: no chunks from matched files — "
                    f"falling back to unfiltered retrieval"
                )
                doc_filter_active = False
                seen_content_hashes.clear()
                for chunk in raw_results:
                    h = _chunk_content_hash(chunk.get("text", ""))
                    if h in seen_content_hashes:
                        continue
                    seen_content_hashes.add(h)
                    result.append(chunk)
                    if len(result) >= k:
                        break
        elif file_boost_map:
            chunk_scores = []
            for rank, chunk in enumerate(raw_results):
                h = _chunk_content_hash(chunk.get("text", ""))
                if h in seen_content_hashes:
                    continue
                seen_content_hashes.add(h)
                sf = chunk.get("source_file", "")
                boost = file_boost_map.get(sf, 1.0)
                rrf_like = 1.0 / (60 + rank + 1)
                chunk_scores.append((chunk, rrf_like * boost))

            chunk_scores.sort(key=lambda x: x[1], reverse=True)

            max_per_group = max(k // 10, 2)
            group_counts = {}
            result_files = set()
            for chunk, score in chunk_scores:
                sf = chunk.get("source_file", "")
                group = self._get_source_group(sf)
                gc = group_counts.get(group, 0)
                if gc < max_per_group:
                    result.append(chunk)
                    result_files.add(sf)
                    group_counts[group] = gc + 1
                if len(result) >= k:
                    break

            # Guaranteed inclusion for top soft-boost files
            if _excel_soft_boost:
                top_boost = sorted(
                    file_boost_map.items(), key=lambda x: x[1], reverse=True
                )[:3]
                for sf, boost_val in top_boost:
                    if sf in result_files:
                        continue
                    if sf not in self._file_chunk_groups:
                        continue
                    bm25_scores = self.bm25.get_scores(query.lower().split())
                    fchunks = self._file_chunk_groups[sf]
                    scored = sorted(
                        [(i, bm25_scores[i]) for i in fchunks],
                        key=lambda x: x[1], reverse=True,
                    )
                    added = 0
                    for ci, _ in scored:
                        if added >= 2:
                            break
                        h = _chunk_content_hash(
                            self.documents[ci].get("text", "")
                        )
                        if h not in seen_content_hashes:
                            seen_content_hashes.add(h)
                            result.insert(0, self.documents[ci])
                            result_files.add(sf)
                            added += 1
                    if added:
                        logger.info(
                            f"  Soft-boost inclusion: added {added} chunks "
                            f"from {sf} (boost={boost_val:.1f})"
                        )
        else:
            for chunk in raw_results:
                h = _chunk_content_hash(chunk.get("text", ""))
                if h in seen_content_hashes:
                    continue
                seen_content_hashes.add(h)
                result.append(chunk)
                if len(result) >= k:
                    break

        # ---- Step 3: Sibling part expansion ----
        sibling_chunks = []
        for chunk in result:
            if chunk.get("section_total_parts", 1) <= 1:
                continue
            heading = chunk.get("heading_text", "")
            source  = chunk.get("source_file", "")
            total   = chunk.get("section_total_parts", 1)
            for doc in self.documents:
                if (doc.get("heading_text") == heading
                        and doc.get("source_file") == source
                        and doc.get("section_total_parts") == total):
                    h = _chunk_content_hash(doc.get("text", ""))
                    if h not in seen_content_hashes:
                        if doc_filter_active and doc.get("source_file", "") not in matched_docs:
                            continue
                        seen_content_hashes.add(h)
                        sibling_chunks.append(doc)

        sibling_chunks.sort(key=lambda c: c.get("section_part", 0))
        result.extend(sibling_chunks)

        # ---- Step 3b: Table expansion ----
        if doc_filter_active:
            tabular_sources_in_result = set()
            for chunk in result:
                if chunk.get("doc_type") == "tabular":
                    tabular_sources_in_result.add(chunk.get("source_file", ""))

            if tabular_sources_in_result:
                table_expansion_chunks = []
                for doc in self.documents:
                    if doc.get("doc_type") != "tabular":
                        continue
                    if doc.get("source_file", "") not in tabular_sources_in_result:
                        continue
                    h = _chunk_content_hash(doc.get("text", ""))
                    if h not in seen_content_hashes:
                        seen_content_hashes.add(h)
                        table_expansion_chunks.append(doc)
                result.extend(table_expansion_chunks)
                if table_expansion_chunks:
                    logger.info(f"  Table expansion: added {len(table_expansion_chunks)} additional table chunks")

        # ---- Step 3c: Hard cap on chunks sent to LLM ----
        MAX_CHUNKS_TO_LLM = 8
        if len(result) > MAX_CHUNKS_TO_LLM:
            logger.info(f"  Chunk cap: trimming {len(result)} -> {MAX_CHUNKS_TO_LLM}")
            result = result[:MAX_CHUNKS_TO_LLM]

        # ---- Step 4: Store retrieval debug info ----
        self._last_retrieval_debug = {
            "query": query,
            "doc_filter_active": doc_filter_active,
            "matched_docs": matched_docs,
            "raw_count": len(raw_results),
            "filtered_count": len(result),
            "sibling_count": len(sibling_chunks),
            "chunks": [
                {
                    "chunk_id": c.get("chunk_id"),
                    "source_file": c.get("source_file", ""),
                    "section": c.get("heading_text", ""),
                    "text": c.get("text", ""),
                    "page": c.get("page_num", ""),
                }
                for c in result
            ],
            "raw_results_preview": [
                {
                    "chunk_id": c.get("chunk_id"),
                    "source": c.get("source_file", "?"),
                    "heading": c.get("heading_text", ""),
                    "text_len": len(c.get("text", "")),
                    "page": c.get("page_num", ""),
                }
                for c in raw_results[:30]
            ],
        }

        logger.info(
            f"  Retrieval: {len(result)} chunks for LLM "
            f"({len(result) - len(sibling_chunks)} from search + {len(sibling_chunks)} sibling parts)"
            f"{' [DOC-FILTERED: ' + str(matched_docs) + ']' if doc_filter_active else ''}"
        )
        return result

    def _retrieve_hybrid(self, query: str, k: int) -> List[Dict[str, Any]]:
        """Hybrid retrieval: FAISS dense + BM25 sparse, fused with RRF."""
        query_embedding = self._get_embedding(query)
        query_vec = np.array([query_embedding]).astype("float32")

        # Dense search (FAISS)
        search_k = min(k, self.vector_store.ntotal)
        D_dense, I_dense = self.vector_store.search(query_vec, search_k)
        dense_results = {int(idx): 1.0 / (i + 60) for i, idx in enumerate(I_dense[0]) if idx != -1}

        # Sparse search (BM25)
        tokenized_query = query.lower().split()
        bm25_scores = self.bm25.get_scores(tokenized_query)
        bm25_ranked_indices = np.argsort(bm25_scores)[::-1][:k]
        sparse_results = {int(idx): 1.0 / (i + 60) for i, idx in enumerate(bm25_ranked_indices)}

        # RRF fusion
        rrf_scores = {}
        for idx in set(list(dense_results.keys()) + list(sparse_results.keys())):
            rrf_scores[idx] = dense_results.get(idx, 0) + sparse_results.get(idx, 0)

        sorted_indices = sorted(rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True)
        return [self.documents[idx] for idx in sorted_indices[:k]]

    # ------------------------------------------------------------------
    # Generation — Ollama Native API with retry logic
    # ------------------------------------------------------------------
    def generate_answer(self, query: str, retrieved_chunks: List[Dict[str, Any]]) -> str:
        """Generate answer using Ollama's native /api/chat endpoint with retry logic."""
        context = ""
        for i, chunk in enumerate(retrieved_chunks):
            source = chunk.get("source_file", "Unknown")
            page = f", Page: {chunk.get('page_num')}" if chunk.get("page_num") else ""
            sheet = f", Sheet: {chunk.get('sheet_name')}" if chunk.get("sheet_name") else ""

            heading = chunk.get("heading_text", "")
            parent_path = chunk.get("parent_path", "")
            hierarchy = ""
            if heading:
                hierarchy = f", Section: {parent_path + ' > ' if parent_path else ''}{heading}"

            # Show part info when a large section was split
            part_info = ""
            if chunk.get("section_total_parts", 1) > 1:
                part_info = f" [Part {chunk['section_part']}/{chunk['section_total_parts']}]"

            context += f"--- Source {i + 1}: {source}{page}{sheet}{hierarchy}{part_info} ---\n{chunk['text']}\n\n"

        messages = [
            {
                "role": "system",
                "content": (
                    "You are a precise assistant that answers questions using ONLY the "
                    "provided context documents.\n\n"
                    "RULES:\n"
                    "1. Answer FULLY and COMPLETELY — do not cut off, summarize too briefly, "
                    "or skip any relevant detail from the context.\n"
                    "2. Each source block contains the COMPLETE text of a section (or a numbered "
                    "part of a large section). Extract EVERY relevant item, sub-point, or "
                    "sub-section it mentions — including 3.1, 3.2, 3.3 ... through the last one.\n"
                    "3. When multiple source blocks are parts of the same section "
                    "(e.g. [Part 1/3], [Part 2/3]), read ALL parts together as one continuous section "
                    "and give a single combined answer.\n"
                    "4. When the context lists items (e.g. reasons, steps, rules, sub-sections), "
                    "include ALL of them — never stop early or say 'and more'.\n"
                    "5. Always cite sources using [Source X] format.\n"
                    "6. If the answer is NOT in the context, say exactly: "
                    "'The provided documents do not contain this information.'\n"
                    "7. Do NOT invent or infer anything not explicitly stated in the context.\n"
                    "8. Use bullet points or numbered lists when the context has multiple items.\n"
                    "9. You will receive chunks from MULTIPLE source documents. You MUST read "
                    "and consider ALL of them — do NOT ignore or skip any source. "
                    "Synthesize information across all relevant sources and cite each one "
                    "using [Source X]. If a source contains relevant information to the question, "
                    "include it in your answer regardless of which document it comes from.\n"
                    "10. When sources from different documents provide overlapping or complementary "
                    "information, combine them into a single coherent answer rather than treating "
                    "each source in isolation.\n"
                    "11. When listing items from a table, you MUST first COUNT every row"
                    "in the source, state the total count found, then list every single"
                    "one. Never stop listing until you have listed ALL rows you counted."

                ),
            },
            {"role": "user", "content": f"Context:\n{context}\n\nQuestion: {query}"},
        ]

        payload = {
            "model": self.model_name,
            "messages": messages,
            "stream": False,
            "options": {
                "temperature": 0,
                "num_predict": OLLAMA_NUM_PREDICT,
                "num_ctx": 16384,
            },
        }

        url = f"{self.ollama_base_url}/api/chat"

        # Retry logic for Ollama
        for attempt in range(1, OLLAMA_MAX_RETRIES + 1):
            try:
                response = requests.post(
                    url, json=payload, timeout=120, verify=self.ssl_verify
                )
                response.raise_for_status()
                result = response.json()
                return result.get("message", {}).get("content", "No response from Ollama.")
            except requests.exceptions.ConnectionError:
                if attempt == OLLAMA_MAX_RETRIES:
                    return (
                        f"Error: Could not connect to Ollama at {self.ollama_base_url}. "
                        f"Make sure Ollama is running: 'ollama serve'"
                    )
                logger.warning(f"Ollama connection retry {attempt}/{OLLAMA_MAX_RETRIES}...")
                time.sleep(OLLAMA_RETRY_DELAY * attempt)
            except requests.exceptions.Timeout:
                if attempt == OLLAMA_MAX_RETRIES:
                    return "Error: Ollama request timed out. Try a smaller model or simpler query."
                logger.warning(f"Ollama timeout retry {attempt}/{OLLAMA_MAX_RETRIES}...")
                time.sleep(OLLAMA_RETRY_DELAY * attempt)
            except Exception as e:
                if attempt == OLLAMA_MAX_RETRIES:
                    return f"Error generating answer: {e}"
                logger.warning(f"Ollama error retry {attempt}/{OLLAMA_MAX_RETRIES}: {e}")
                time.sleep(OLLAMA_RETRY_DELAY * attempt)

        return "Error: All Ollama retries exhausted."

    # ------------------------------------------------------------------
    # Query Reformulation — make vague queries better for retrieval
    # ------------------------------------------------------------------
    def _reformulate_query(self, query: str) -> str:
        """
        Reformulate a user's vague query into a better search query.

        Uses the LLM to expand short/vague queries into more specific
        retrieval-friendly queries. For example:
          "section 3" -> "section 3 complete details all subsections 3.1 3.2 3.3"
          "what is the policy" -> "policy definition scope applicability requirements"

        If the query is already specific (>20 chars with detail words),
        it passes through unchanged.
        """
        # Skip reformulation for already-detailed queries
        detail_indicators = [
            "explain", "describe", "details", "including", "specifically",
            "comprehensive", "elaborate", "full", "complete", "all about",
            "summary of", "overview of", "definition of", "how does",
            "what are the", "list all", "compare", "difference between",
        ]
        query_lower = query.lower()
        if len(query.split()) > 8 or any(d in query_lower for d in detail_indicators):
            logger.info("  Query is already specific — skipping reformulation")
            return query

        # Use LLM to reformulate
        reformulate_prompt = (
            "You are a query reformulation assistant for a document search system. "
            "Your job is to take a short or vague user query and expand it into a "
            "more detailed search query that will help find ALL relevant information.\n\n"
            "Rules:\n"
            "1. Keep the original intent — don't change what the user is asking\n"
            "2. Add synonyms and related terms that might appear in documents\n"
            "3. If the query mentions a section/heading, add words like 'details', "
            "'subsections', 'full content', 'all parts'\n"
            "4. Keep it concise — 1-2 sentences max\n"
            "5. Return ONLY the reformulated query, nothing else\n\n"
            f"User query: {query}\n\nReformulated query:"
        )

        payload = {
            "model": self.model_name,
            "messages": [
                {"role": "user", "content": reformulate_prompt}
            ],
            "stream": False,
            "options": {"temperature": 0.3},
        }

        url = f"{self.ollama_base_url}/api/chat"

        try:
            response = requests.post(
                url, json=payload, timeout=30, verify=self.ssl_verify
            )
            response.raise_for_status()
            result = response.json()
            reformulated = result.get("message", {}).get("content", "").strip()

            if reformulated and len(reformulated) > len(query):
                logger.info(f"  Query reformulated: '{query}' -> '{reformulated}'")
                return reformulated
            else:
                logger.info("  Reformulation not better — using original query")
                return query

        except Exception as e:
            logger.warning(f"  Query reformulation failed: {e} — using original query")
            return query

    # ------------------------------------------------------------------
    # Debug — Save chunks passed to LLM as a text file
    # ------------------------------------------------------------------
    def _save_debug_chunks(
        self,
        query: str,
        chunks: List[Dict[str, Any]],
        search_query: str = "",
    ):
        """
        Save the chunks that are being passed to the LLM as a detailed
        text file for debugging retrieval and context construction.

        This file shows:
          - The original user query and the reformulated search query
          - Every chunk with full metadata (role, section, parent, text)
          - Page coverage summary — which pages are represented
          - Summary stats at the bottom

        The file is APPENDED to (not overwritten) so you can see the
        history of multiple queries in one session.
        """
        if not self.debug_chunks_file:
            return

        try:
            lines = []
            lines.append("=" * 80)
            lines.append(f"  DEBUG: Chunks Passed to LLM")
            lines.append(f"  Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append("=" * 80)
            lines.append("")
            lines.append(f"  User Query:        {query}")
            if search_query and search_query != query:
                lines.append(f"  Search Query:      {search_query}  (reformulated)")
            else:
                lines.append(f"  Search Query:      {query}  (original, no reformulation)")
            lines.append(f"  Total Chunks:      {len(chunks)}")
            lines.append("")

            # ---- NEW: Document filter info ----
            debug_info = getattr(self, '_last_retrieval_debug', None)
            if debug_info:
                lines.append(f"  DOC FILTER:        {'ACTIVE' if debug_info['doc_filter_active'] else 'OFF'}")
                if debug_info['doc_filter_active']:
                    for md in debug_info['matched_docs']:
                        # Show the clean name alongside the raw filename
                        raw_stem = os.path.splitext(os.path.basename(md))[0].lower()
                        clean_name = self._get_meaningful_name(raw_stem)
                        if clean_name != raw_stem:
                            lines.append(f"    -> Matched: {md}  (clean name: '{clean_name}')")
                        else:
                            lines.append(f"    -> Matched: {md}")
                elif debug_info.get('generic_guard_triggered'):
                    lines.append(f"    -> GENERIC_GUARD: query too generic, matched {debug_info['generic_guard_count']} different docs")
                lines.append(f"  Raw search hits:   {debug_info['raw_count']}")
                lines.append(f"  After filter:      {debug_info['filtered_count']}")
                lines.append(f"  Sibling parts:     {debug_info['sibling_count']}")
                lines.append("")

                # ---- NEW: Raw top-k results table ----
                lines.append("-" * 80)
                lines.append("  RAW TOP-K RESULTS (before doc filter & dedup)")
                lines.append("-" * 80)
                for i, r in enumerate(debug_info.get("raw_results_preview", [])):
                    marker = ""
                    if debug_info['doc_filter_active']:
                        marker = " ✓" if r["source"] in debug_info['matched_docs'] else " ✗ (filtered out)"
                    lines.append(f"  {i+1:2d}. [ID={r['chunk_id']}] {r['source']} | {r['heading'] or '(no heading)'} | p{r['page']} | {r['text_len']}c{marker}")
                lines.append("")

            # Role breakdown
            role_counts = {}
            for c in chunks:
                role = c.get("chunk_role", "standalone")
                role_counts[role] = role_counts.get(role, 0) + 1
            lines.append(f"  Role Breakdown:")
            for role, count in role_counts.items():
                lines.append(f"    {role}: {count}")
            lines.append("")

            # Source file breakdown
            source_counts = {}
            for c in chunks:
                src = c.get("source_file", "Unknown")
                source_counts[src] = source_counts.get(src, 0) + 1
            lines.append(f"  Source Breakdown:")
            for src, count in source_counts.items():
                lines.append(f"    {src}: {count}")
            lines.append("")

            # Page coverage summary — shows which pages are represented
            page_coverage = {}
            for c in chunks:
                page = c.get("page_num")
                if page:
                    src = c.get("source_file", "Unknown")
                    key = f"{src} page {page}"
                    page_coverage[key] = page_coverage.get(key, 0) + 1
            if page_coverage:
                lines.append(f"  Page Coverage (pages present in these chunks):")
                for key, count in sorted(page_coverage.items()):
                    lines.append(f"    {key}: {count} chunk(s)")
                lines.append(f"  WARNING: If a page is NOT listed above, it was NOT retrieved!")
                lines.append("")

            # Detailed chunk listing
            lines.append("-" * 80)
            lines.append("  DETAILED CHUNK LISTING")
            lines.append("-" * 80)

            for i, chunk in enumerate(chunks):
                chunk_id    = chunk.get("chunk_id", "?")
                role        = chunk.get("chunk_role", "standalone")
                source      = chunk.get("source_file", "?")
                heading     = chunk.get("heading_text", "")
                parent_path = chunk.get("parent_path", "")
                hlevel      = chunk.get("heading_level", 0)
                parent_id   = chunk.get("parent_id")
                child_ids   = chunk.get("child_ids", [])
                page_num    = chunk.get("page_num", "")
                doc_type    = chunk.get("doc_type", "narrative")
                text        = chunk.get("text", "")
                text_len    = len(text)

                role_label = {
                    "parent": "PARENT (section overview)",
                    "child": "CHILD (subsection)",
                    "standalone": "STANDALONE",
                }.get(role, role.upper())

                lines.append("")
                lines.append(f"  --- Chunk #{i + 1} [ID={chunk_id}] ---")
                lines.append(f"  Role:         {role_label}")
                lines.append(f"  Source:       {source}")
                if page_num:
                    lines.append(f"  Page:         {page_num}")
                lines.append(f"  Doc Type:     {doc_type}")
                lines.append(f"  Text Length:  {text_len} chars")

                if heading:
                    hierarchy = f"{parent_path} > {heading}" if parent_path else heading
                    lines.append(f"  Section:      {hierarchy} (H{hlevel})")
                if parent_path and not heading:
                    lines.append(f"  Parent Path:  {parent_path}")
                if role == "parent" and child_ids:
                    lines.append(f"  Child IDs:    {child_ids}")
                if role == "child" and parent_id is not None:
                    lines.append(f"  Parent ID:    {parent_id}")

                lines.append(f"  Text:")
                lines.append(f"  {text}")
                lines.append(f"  --- End Chunk #{i + 1} ---")

            # Summary footer
            lines.append("")
            lines.append("=" * 80)
            total_chars = sum(len(c.get("text", "")) for c in chunks)
            lines.append(f"  TOTAL: {len(chunks)} chunks, {total_chars} total characters")
            lines.append(f"  Approx tokens: ~{total_chars // 4} (rough estimate)")
            lines.append("=" * 80)
            lines.append("")
            lines.append("")

            # Append to file (not overwrite — so multiple queries accumulate)
            with open(self.debug_chunks_file, "a", encoding="utf-8") as f:
                f.write("\n".join(lines))

            logger.info(
                f"  Debug chunks saved to: {os.path.abspath(self.debug_chunks_file)} "
                f"({len(chunks)} chunks, {total_chars} chars)"
            )

        except Exception as e:
            logger.warning(f"  Failed to save debug chunks file: {e}")

    # ------------------------------------------------------------------
    # Chat entry point
    # ------------------------------------------------------------------
    def chat(self, query: str, source_type: str = "all") -> str:
        search_query = query

        # Step 2: Retrieve using the original query
        chunks = self.retrieve(search_query, source_type=source_type)
        if not chunks:
            return "No relevant documents found. Please ingest documents first."

        # Step 3: Save debug chunks to text file for inspection
        self._save_debug_chunks(query, chunks, search_query)

        # Step 4: Generate answer
        return self.generate_answer(query, chunks)

    # ------------------------------------------------------------------
    # Reset — clear index and manifest for fresh start
    # ------------------------------------------------------------------
    def reset(self):
        """Reset the system — clear index, documents, and manifest."""
        self.vector_store = None
        self._index_type = "flat"
        self.documents = []
        self.bm25 = None
        self._file_bm25 = None
        self._file_names_ordered = []
        self._file_chunk_groups = {}
        self._chunk_hashes: set = set()
        self._last_retrieval_debug = None
        self._ingestion_stats = {
            "total_docs": 0, "total_chunks": 0,
            "failed_docs": 0, "start_time": None, "end_time": None,
        }
        self._failed_doc_details: List[Tuple[str, str]] = []

        # Clear manifest file
        if os.path.exists(self.manifest.manifest_path):
            os.remove(self.manifest.manifest_path)
        self.manifest = IngestionManifest(self.manifest.manifest_path)

        # Clear local FAISS files
        if os.path.exists(self.faiss_local_path):
            import shutil
            shutil.rmtree(self.faiss_local_path, ignore_errors=True)
            logger.info(f"Local index cleared: {self.faiss_local_path}")

        logger.info("RAG system reset complete. Run --ingest to start fresh.")